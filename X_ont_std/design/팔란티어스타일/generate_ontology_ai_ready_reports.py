from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
TODAY = date(2026, 6, 8).strftime("%Y-%m-%d")


PALETTE = {
    "ink": RGBColor(31, 41, 55),
    "muted": RGBColor(91, 103, 120),
    "navy": RGBColor(21, 46, 87),
    "blue": RGBColor(37, 99, 235),
    "teal": RGBColor(13, 148, 136),
    "slate": RGBColor(71, 85, 105),
    "light_blue": "EAF2FF",
    "light_teal": "E6F7F4",
    "light_slate": "F3F6FA",
    "border": "D8DEE9",
    "navy_fill": "152E57",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D8DEE9", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_widths(table, widths_cm) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def style_run(run, bold=False, size=None, color=None) -> None:
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.font.color.rgb = PALETTE["ink"]
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in (
        ("Title", 20, PALETTE["navy"]),
        ("Heading 1", 15, PALETTE["navy"]),
        ("Heading 2", 12, PALETTE["blue"]),
        ("Heading 3", 10.5, PALETTE["teal"]),
    ):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 7)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    style_run(run, bold=True, size=21, color=PALETTE["navy"])

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    style_run(run, size=10.5, color=PALETTE["muted"])

    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_col_widths(table, [3.2, 13.6])
    for row_idx, (label, value) in enumerate(meta):
        cells = table.rows[row_idx].cells
        cells[0].text = label
        cells[1].text = value
        for c in cells:
            set_cell_border(c)
            set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], PALETTE["light_slate"])
        for run in cells[0].paragraphs[0].runs:
            style_run(run, bold=True, size=9.5, color=PALETTE["slate"])
        for run in cells[1].paragraphs[0].runs:
            style_run(run, size=9.5, color=PALETTE["ink"])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = "EAF2FF") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "C7D2FE")
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    style_run(r, bold=True, size=10.5, color=PALETTE["navy"])
    p = cell.add_paragraph()
    r = p.add_run(body)
    style_run(r, size=9.5, color=PALETTE["ink"])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        style_run(run, size=9.7, color=PALETTE["ink"])


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        style_run(run, size=9.7, color=PALETTE["ink"])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_col_widths(table, widths_cm)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = header
        set_cell_shading(cell, PALETTE["navy_fill"])
        set_cell_border(cell, "B8C2D6")
        set_cell_margins(cell, 120, 120, 120, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                style_run(run, bold=True, size=9, color=RGBColor(255, 255, 255))
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx], 110, 120, 110, 120)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 0:
                set_cell_shading(cells[idx], "FBFCFE")
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    style_run(run, size=8.7, color=PALETTE["ink"])
    doc.add_paragraph()


def add_sources(doc: Document, sources: list[tuple[str, str]]) -> None:
    doc.add_heading("참고 공개 자료", level=1)
    for name, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.45)
        run = p.add_run(f"{name}: {url}")
        style_run(run, size=8.8, color=PALETTE["muted"])


def build_ontology_workflow_report() -> Path:
    doc = Document()
    configure_doc(doc)
    add_title_block(
        doc,
        "온톨로지+워크플로우 고도화 테크니컬 리포트",
        "엔코아 협업 전제의 한국형 엔터프라이즈 온톨로지 플랫폼 기술 방향",
        [
            ("작성일", TODAY),
            ("작성 도구", "Codex"),
            ("범위", "온톨로지 메타모델, DB 매핑, 쿼리 엔진, 정확도/성능 검증, 워크플로우 확장"),
            ("전제", "엔코아급 데이터 모델링/튜닝 역량과 결합하여 6개월~2년 로드맵으로 추진"),
            ("포지션", "BI/로우코드 중심이 아닌 데이터 모델링 기반 온톨로지 코어 우위 확보"),
        ],
    )

    doc.add_heading("1. 핵심 판단", level=1)
    add_callout(
        doc,
        "전략 결론",
        "현재 프로토타입 관점에서는 JSONB 기반 SPARQL 변환기로 충분하지만, 엔터프라이즈 제품 관점에서는 ROW형 canonical ontology schema, DB 매핑 계층, 정확도/성능 벤치마크, 변경 영향도 관리가 핵심 경쟁력이 된다. 워크플로우 자동화는 온톨로지 코어가 안정된 뒤 2순위로 확장한다.",
    )
    add_bullets(
        doc,
        [
            "직접적인 팔란티어 벤치마크는 불가능하므로 공개 문서 기반의 아키텍처 추론 비교와 자체 벤치마크 체계를 분리한다.",
            "BI매트릭스 대비 우위는 화면/BI 편의성이 아니라 DB 모델링, 정합성, 성능 튜닝, 온톨로지-물리모델 매핑에서 만들어야 한다.",
            "엔코아의 데이터 표준화, 모델링, 품질, 튜닝 역량은 온톨로지 플랫폼의 진입장벽을 만들 수 있는 핵심 자산이다.",
        ],
    )

    doc.add_heading("2. 목표 아키텍처", level=1)
    add_table(
        doc,
        ["계층", "핵심 역할", "기술 포인트", "우선순위"],
        [
            ["Ontology Core", "업무 개념, 속성, 관계, 제약의 canonical model", "ROW형 메타모델 + JSONB 확장 속성 병행", "1순위"],
            ["Source Mapping", "물리 DB/컬럼/API/문서와 온톨로지 연결", "표준용어, ERD, PK/FK, lineage 수집", "1순위"],
            ["Query Engine", "SPARQL-like, SQL, graph traversal 질의", "핫패스 SQL 변환, 인덱스, 캐시, 쿼리 플래너", "1순위"],
            ["Accuracy Layer", "AI 답변과 온톨로지 근거의 일치 검증", "EvidenceGate, golden set, 실패 유형 분석", "1순위"],
            ["Workflow Layer", "업무 action, 승인, write-back, 알림", "Action type, command handler, transaction boundary", "2순위"],
            ["Governance", "권한, 감사, 버전, 영향도 관리", "객체/관계 단위 정책, 변경 이력, 영향도 그래프", "1.5순위"],
        ],
        [3.0, 5.0, 6.1, 2.0],
    )

    doc.add_heading("3. ROW형 Canonical Ontology Schema", level=1)
    add_table(
        doc,
        ["테이블/개념", "목적", "핵심 컬럼 예시"],
        [
            ["ontology_class", "업무 객체/개념 타입 정의", "class_id, name, domain, lifecycle, owner, version"],
            ["ontology_property", "속성 정의와 타입/제약 관리", "property_id, class_id, data_type, nullable, validation_rule"],
            ["ontology_relation", "관계 타입과 방향성/cardinality 정의", "relation_id, from_class, to_class, relation_type, cardinality"],
            ["entity_instance", "실제 인스턴스 저장", "entity_id, class_id, source_id, quality_status"],
            ["property_value", "자주 질의되는 속성값의 정규화 저장", "entity_id, property_id, value_typed, effective_date"],
            ["relation_instance", "실제 관계 저장 및 탐색 최적화", "from_entity, to_entity, relation_id, confidence"],
            ["source_mapping", "물리 데이터와 온톨로지의 연결", "table_name, column_name, class_id, property_id, transform_rule"],
            ["lineage_event", "출처, 변경, 변환 근거", "event_id, source, operation, actor, timestamp"],
        ],
        [4.0, 5.2, 7.0],
    )

    doc.add_heading("4. 성능/정확도 우선 검증 체계", level=1)
    add_table(
        doc,
        ["검증 축", "측정 지표", "초기 목표", "비고"],
        [
            ["쿼리 성능", "p50/p95 응답시간, QPS, 캐시 hit rate", "100K entity / 1M relation에서 주요 핫패스 p95 200ms 이하", "직접 측정값만 보고서에 사용"],
            ["정확도", "answer precision, evidence recall, unsupported answer rate", "도메인 질문셋 기준 90% 이상", "RAG 답변과 온톨로지 근거를 분리 평가"],
            ["정합성", "타입 위반, cardinality 위반, 중복 개념, 매핑 누락", "릴리즈 전 자동 검증 100%", "품질 게이트화"],
            ["변경 안정성", "영향도 분석 커버리지, 회귀 테스트 통과율", "주요 API/쿼리/AI 답변 영향도 자동 산출", "엔터프라이즈 필수"],
        ],
        [3.0, 5.6, 4.0, 4.2],
    )

    doc.add_heading("5. 추론적 경쟁 비교", level=1)
    add_table(
        doc,
        ["비교 항목", "자사 방향", "팔란티어 추론", "BI매트릭스 추론", "우위 전략"],
        [
            ["온톨로지 코어", "DB 모델링과 메타데이터 기반으로 정교화", "운영 계층까지 통합된 성숙 플랫폼", "BI/AI 분석 UX 중심 가능성", "한국 기업 DB 현실에 밀착한 모델링 표준"],
            ["성능", "RDB 튜닝/인덱싱/핫패스 최적화", "고속 인덱싱/검색 백엔드 보유 가능성", "DB 튜닝 전문성은 공개 정보상 약점 가능", "엔코아 튜닝 역량을 제품 기능화"],
            ["정확도", "EvidenceGate와 온톨로지 검증으로 통제", "거버넌스/액션/보안까지 연결", "생성형 AI 분석 정확도 중심 가능", "근거 기반 답변/변경 영향도 차별화"],
            ["워크플로우", "2단계 확장", "Action/Function/Write-back 통합", "로우코드 업무 앱 강점 가능", "온톨로지 코어 안정 후 action type 도입"],
        ],
        [3.1, 4.4, 4.4, 4.3, 4.2],
    )

    doc.add_heading("6. 6개월~2년 로드맵", level=1)
    add_table(
        doc,
        ["기간", "목표", "주요 산출물", "성공 기준"],
        [
            ["0~6개월", "온톨로지 코어 MVP", "ROW형 메타모델, source mapping, 핫패스 쿼리, 정확도 평가셋", "100K/1M 기준 벤치마크와 100문항 정확도 리포트"],
            ["6~12개월", "엔터프라이즈 온톨로지 플랫폼 v1", "Ontology Studio, 영향도 분석, 품질 게이트, 하이브리드 검색", "파일럿 고객 1~2곳에서 실제 DB 매핑 검증"],
            ["12~24개월", "한국형 팔란티어 경량 버전", "Action type, workflow, write-back governance, 산업 템플릿", "금융/공공/제조 중 1개 산업 레퍼런스 확보"],
        ],
        [2.4, 4.3, 6.5, 4.8],
    )

    doc.add_heading("7. 기술 리스크와 대응", level=1)
    add_table(
        doc,
        ["리스크", "영향", "대응"],
        [
            ["온톨로지 과설계", "초기 개발 지연과 사용자 이탈", "MVP는 class/property/relation/source_mapping으로 제한"],
            ["SPARQL 완전 지원 과장", "신뢰도 하락", "지원 범위를 명시하고 W3C 완전 구현과 구분"],
            ["AI 답변 환각", "엔터프라이즈 도입 실패", "EvidenceGate, 근거 없는 답변 차단, 평가셋 운영"],
            ["워크플로우 선행 구현", "코어 정합성 약화", "Action type은 12개월 이후 확장"],
            ["경쟁사 대비 UX 열세", "PoC 설득력 약화", "기술 우위 자료와 최소한의 Studio UX 병행"],
        ],
        [4.0, 5.0, 7.0],
    )

    add_sources(
        doc,
        [
            ("Palantir Ontology overview", "https://www.palantir.com/docs/foundry/ontology/overview"),
            ("Palantir Ontology query compute", "https://www.palantir.com/docs/foundry/ontologies/query-compute-usage"),
            ("W3C SPARQL 1.1 Query Language", "https://www.w3.org/TR/sparql11-query/"),
            ("엔코아 데이터 컨설팅/솔루션 공개 소개", "https://www.sharedit.co.kr/partners/1271"),
            ("BI매트릭스 TRINITY 온프레미스 AI 보도", "https://zdnet.co.kr/view/?no=20260423180720"),
        ],
    )
    path = OUT_DIR / "Codex_온톨로지_워크플로우_고도화_테크니컬_리포트.docx"
    doc.save(path)
    return path


def build_ai_ready_report() -> Path:
    doc = Document()
    configure_doc(doc)
    add_title_block(
        doc,
        "토탈 AI-Ready 확장 방안 보고서",
        "부트캠프, 개발 표준, 데이터 모델링 표준, 튜닝 가이드, AI 코딩 에이전트 연동 전략",
        [
            ("작성일", TODAY),
            ("작성 도구", "Codex"),
            ("범위", "AI 백엔드 개발 표준, 엔코아 부트캠프, 표준 가이드, 온톨로지 기반 코딩 에이전트"),
            ("전제", "OCP/튜닝교육/금융권 대량 데이터/DB 전공 경험을 토대로 교육-표준-플랫폼을 연결"),
            ("목표", "AI가 기업 백엔드를 안전하게 개발/변경/검증할 수 있는 표준 체계 구축"),
        ],
    )

    doc.add_heading("1. 핵심 방향", level=1)
    add_callout(
        doc,
        "전략 결론",
        "AI-Ready 확장은 온톨로지를 지식 그래프에 한정하지 않고, DB 모델, API, 테스트, 보안, 튜닝, 코드 리뷰, AI 코딩 에이전트의 공통 기준 언어로 확장하는 전략이다. 교육은 단기 매출과 인력 양성 채널이면서, 장기 플랫폼의 표준 검증장 역할을 한다.",
        fill=PALETTE["light_teal"],
    )
    add_bullets(
        doc,
        [
            "부트캠프는 개발 표준 가이드 v1을 실전 검증하는 채널로 설계한다.",
            "개발 표준은 문서에서 끝내지 않고 템플릿, 코드 생성기, 품질 게이트, AI agent context API로 전환한다.",
            "데이터 모델링 표준과 튜닝 가이드는 엔코아 협업 시 가장 강한 차별화 자산이 된다.",
            "AI 코딩 에이전트는 자율 개발자가 아니라 표준을 따르는 controlled implementer로 포지셔닝한다.",
        ],
    )

    doc.add_heading("2. 제품/사업 확장 구조", level=1)
    add_table(
        doc,
        ["단계", "산출물", "역할", "확장 포인트"],
        [
            ["교육", "엔코아 부트캠프", "표준 기반 AI 백엔드 개발자 양성", "수강생/고객 피드백으로 표준 검증"],
            ["문서", "개발 표준 가이드", "API/서비스/DB/테스트/보안 기준", "기업별 커스터마이징"],
            ["자동화", "표준 템플릿/검사기", "코드 생성과 리뷰 자동화", "CI/CD 품질 게이트"],
            ["모델링", "데이터 모델링 표준 가이드", "온톨로지와 RDB 모델 연결", "엔코아 방법론 연계"],
            ["튜닝", "AI-ready 튜닝 가이드", "대용량 쿼리와 RAG 검색 성능 확보", "금융/공공 PoC 대응"],
            ["플랫폼", "AI Coding Agent Context Server", "AI 에이전트가 표준과 온톨로지를 조회", "엔터프라이즈 개발 운영 플랫폼"],
        ],
        [2.5, 4.5, 5.0, 5.1],
    )

    doc.add_heading("3. AI 백엔드 개발 표준 범위", level=1)
    add_table(
        doc,
        ["표준 영역", "규정할 내용", "AI 에이전트 연동"],
        [
            ["API 설계", "resource naming, endpoint contract, pagination, error schema", "API scaffold와 contract test 생성"],
            ["도메인 모델", "entity/value object/service/repository 구분", "온톨로지 class/property 기반 코드 골격 생성"],
            ["DB 변경", "migration, rollback, FK/index, 이력 테이블 기준", "변경 영향도 분석 후 migration draft 생성"],
            ["보안/권한", "role, data scope, row-level access, audit", "권한 누락 코드 리뷰와 guard 생성"],
            ["테스트", "unit/integration/e2e/performance test 최소 기준", "테스트 케이스 자동 제안"],
            ["운영성", "logging, tracing, metrics, failure handling", "관측성 표준 누락 탐지"],
        ],
        [3.6, 6.2, 6.8],
    )

    doc.add_heading("4. 데이터 모델링 표준 확장", level=1)
    add_bullets(
        doc,
        [
            "표준 용어, 엔터티, 속성, 관계, 식별자, 코드성 데이터, 이력 모델링 기준을 온톨로지 메타모델과 연결한다.",
            "RDB 테이블/컬럼/관계를 ontology_class, ontology_property, ontology_relation으로 매핑하는 규칙을 표준화한다.",
            "모델 변경 시 API, 쿼리, AI 답변, 워크플로우에 미치는 영향을 자동 산출하는 영향도 분석 체계를 만든다.",
            "엔코아의 데이터 품질/거버넌스 경험을 온톨로지 품질 게이트로 제품화한다.",
        ],
    )

    doc.add_heading("5. 튜닝 가이드 확장", level=1)
    add_table(
        doc,
        ["튜닝 영역", "핵심 주제", "제품화 방식"],
        [
            ["RDB 튜닝", "실행계획, 인덱스, 조인 순서, 통계정보", "튜닝 체크리스트와 query advisor"],
            ["온톨로지 쿼리", "1-hop/2-hop/3-hop 탐색, relation index, materialized path", "핫패스 벤치마크와 SQL rewrite"],
            ["JSONB/ROW 하이브리드", "확장성 vs 인덱싱 성능 선택 기준", "projection table 추천"],
            ["RAG 검색", "vector + keyword + graph hybrid retrieval", "근거 recall/latency 동시 측정"],
            ["캐싱", "query result cache, embedding cache, semantic cache", "캐시 무효화 정책과 hit-rate 대시보드"],
        ],
        [3.4, 6.2, 6.6],
    )

    doc.add_heading("6. AI Coding Agent Context Server", level=1)
    add_table(
        doc,
        ["기능", "설명", "필수 API 예시"],
        [
            ["표준 조회", "개발 표준, 금지 패턴, 필수 테스트 기준 제공", "GET /standards/backend"],
            ["온톨로지 조회", "class/property/relation/action metadata 제공", "GET /ontology/classes/{id}"],
            ["코드 컨텍스트", "기존 서비스/DTO/repository 구조 검색", "GET /code/context?domain=..."],
            ["영향도 분석", "모델 변경이 API/DB/테스트/문서에 미치는 영향 산출", "POST /impact/analyze"],
            ["작업지시 생성", "AI 에이전트가 수행할 안전한 task spec 생성", "POST /agent/tasks"],
            ["검증", "표준 위반, 테스트 누락, 보안 누락 검사", "POST /agent/validate"],
        ],
        [3.4, 6.2, 6.6],
    )

    doc.add_heading("7. 6개월~2년 실행 로드맵", level=1)
    add_table(
        doc,
        ["기간", "중점", "주요 산출물", "판정 기준"],
        [
            ["0~3개월", "표준 v1과 교육 MVP", "부트캠프 커리큘럼, 개발 표준 가이드, 예제 프로젝트", "교육/PoC에서 반복 가능한 개발 패턴 확보"],
            ["3~6개월", "표준 자동화", "코드 템플릿, API scaffold, 테스트 scaffold, agent context API v1", "AI 생성 코드의 표준 준수율 측정"],
            ["6~12개월", "모델링/온톨로지 결합", "데이터 모델링 표준, 온톨로지 매핑, 영향도 분석", "실제 기업 DB 1건 이상 매핑 검증"],
            ["12~24개월", "토탈 AI-ready 플랫폼", "튜닝 advisor, workflow/action, governance, CI/CD 연동", "금융/공공/제조 PoC 레퍼런스 확보"],
        ],
        [2.6, 4.3, 6.4, 4.8],
    )

    doc.add_heading("8. 우선순위", level=1)
    add_numbered(
        doc,
        [
            "개발 표준 가이드 v1을 먼저 만든다. 교육, 코드 생성, 에이전트 연동의 기준점이 된다.",
            "온톨로지 canonical schema와 DB 매핑 표준을 병행 설계한다. 이것이 장기 제품의 중심축이다.",
            "AI 코딩 에이전트는 초기에는 scaffold와 review assistant로 제한한다.",
            "튜닝 가이드는 6개월 이후 실제 벤치마크와 함께 확장한다.",
            "워크플로우/action type은 온톨로지와 개발 표준이 안정된 뒤 12개월 이후 본격화한다.",
        ],
    )

    add_sources(
        doc,
        [
            ("Palantir Ontology overview", "https://www.palantir.com/docs/foundry/ontology/overview"),
            ("Palantir Action types overview", "https://www.palantir.com/docs/foundry/action-types/overview"),
            ("W3C SPARQL 1.1 Query Language", "https://www.w3.org/TR/sparql11-query/"),
            ("엔코아 데이터 컨설팅/솔루션 공개 소개", "https://www.sharedit.co.kr/partners/1271"),
            ("BI매트릭스 TRINITY 온프레미스 AI 보도", "https://zdnet.co.kr/view/?no=20260423180720"),
        ],
    )
    path = OUT_DIR / "Codex_토탈_AI_Ready_확장_방안_보고서.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [build_ontology_workflow_report(), build_ai_ready_report()]
    for path in paths:
        print(path)
