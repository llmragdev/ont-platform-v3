from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUT_PATH = BASE_DIR / "온톨로지_워크플로우_RAG_통합솔루션_시연설명_매뉴얼.docx"


PALETTE = {
    "navy": "0F172A",
    "teal": "0F766E",
    "teal_light": "CCFBF1",
    "slate": "334155",
    "muted": "64748B",
    "line": "CBD5E1",
    "soft": "F8FAFC",
    "blue_soft": "EFF6FF",
    "green_soft": "ECFDF5",
    "amber_soft": "FFFBEB",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = PALETTE["line"], size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None, name: str = "맑은 고딕") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=PALETTE["navy"])
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_para(doc: Document, text: str = "", *, bold_prefix: str | None = None, color: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10, bold=True, color=color or PALETTE["slate"])
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=10, color=color or PALETTE["slate"])
    else:
        r = p.add_run(text)
        set_font(r, size=10, color=color or PALETTE["slate"])
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_font(run, size=10, color=PALETTE["slate"])
        if not p.runs:
            run = p.add_run(item)
            set_font(run, size=10, color=PALETTE["slate"])
        else:
            p.runs[0].text = item


def add_callout(doc: Document, title: str, body: str, fill: str = PALETTE["blue_soft"]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, "BFDBFE")
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=PALETTE["teal"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(body)
    set_font(r2, size=10, color=PALETTE["slate"])
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], PALETTE["teal"])
        set_cell_borders(hdr_cells[i], "0F766E")
        set_cell_margins(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, size=9, bold=True, color="FFFFFF")
    for row in rows:
        new_row = table.add_row()
        tr_pr = new_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        cells = new_row.cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_borders(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                set_cell_shading(cells[i], PALETTE["soft"])
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 12 and "\n" not in value else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_font(run, size=8.5, bold=(i == 0), color=PALETTE["slate"])
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)

    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Ontology + Workflow + RAG 통합 솔루션 시연 설명 매뉴얼")
    set_font(r, size=8, color=PALETTE["muted"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Codex 작성본 | 내부 시연 및 외부 설명 자료")
    set_font(r, size=8, color=PALETTE["muted"])


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Ontology Console v5")
    set_font(r, size=14, bold=True, color=PALETTE["teal"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("온톨로지 + 워크플로우 + RAG\n통합 솔루션 시연 설명 매뉴얼")
    set_font(r, size=28, bold=True, color=PALETTE["navy"])
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    r = p.add_run("시연 후 설명, 메뉴 소개, 공장/고객사 시나리오, 향후 기능 개선 방향")
    set_font(r, size=12, color=PALETTE["slate"])
    p.paragraph_format.space_after = Pt(24)

    add_callout(
        doc,
        "문서 성격",
        "이 문서는 제가 작성한 10~17번 매뉴얼을 바탕으로 만든 시연 설명용 통합 Word 문서입니다. "
        "외부 데모 직후 화면을 설명하거나, 내부 공유용으로 솔루션의 구조와 가치를 설명할 때 사용합니다.",
        fill=PALETTE["green_soft"],
    )

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["작성 기준", "design/메뉴얼 10~17번 Codex 작성 매뉴얼"],
            ["작성일", date.today().isoformat()],
            ["대상 독자", "투자사, 포트폴리오 기업, 현업 담당자, 기술 검토자"],
            ["핵심 메시지", "AI 챗봇이 아니라 업무 실행형 AI 플랫폼 PoC"],
        ],
        widths=[1.6, 4.9],
    )
    doc.add_page_break()


def executive_summary(doc: Document) -> None:
    add_heading(doc, "1. 핵심 요약", 1)
    add_callout(
        doc,
        "한 문장 설명",
        "이 솔루션은 RAG로 근거를 찾고, 워크플로우로 업무를 실행하며, 온톨로지로 처리 결과의 관계를 남기는 업무 실행형 AI 플랫폼입니다.",
    )
    add_table(
        doc,
        ["구성요소", "역할", "시연 시 설명 포인트"],
        [
            ["워크플로우", "업무 처리 절차를 블록으로 구성하고 실행", "어떤 단계에서 어떤 판단과 실행이 일어났는지 보입니다."],
            ["RAG", "문서와 지식 근거 검색", "답변이나 조치가 문서 근거와 연결됩니다."],
            ["온톨로지", "업무 객체와 관계 저장", "요청, 설비, 답변, 정비지시가 관계로 남습니다."],
            ["MCP 연동", "외부 시스템 등록 중계", "댓글/정비지시가 실제 모의 시스템에 등록됩니다."],
            ["스킬", "워크플로우에서 재사용하는 실행 기능", "고객사별 외부 기능을 표준화해 붙일 수 있습니다."],
        ],
        widths=[1.4, 2.1, 3.0],
    )
    add_heading(doc, "1.1 시연에서 강조할 차별점", 2)
    add_bullets(
        doc,
        [
            "단순 RAG 답변이 아니라 외부 게시판 댓글과 정비지시까지 실행합니다.",
            "워크플로우 실행 상태, 입출력, 실패 이력을 추적할 수 있습니다.",
            "실행 결과가 온톨로지 객체와 관계로 저장되어 나중에 분석할 수 있습니다.",
            "company_id, project_id 기준으로 고객사/프로젝트 분리 구조를 설명할 수 있습니다.",
            "스킬 구조를 통해 외부 시스템 연동을 표준화하고 재사용할 수 있습니다.",
        ],
    )


def demo_flow(doc: Document) -> None:
    add_heading(doc, "2. 권장 시연 흐름", 1)
    add_table(
        doc,
        ["순서", "화면", "설명 내용", "성공 확인"],
        [
            ["1", "공장 모의 게시판", "현장 고장/품질 요청이 접수된 상황을 보여줍니다.", "open 요청 존재"],
            ["2", "빌더와 실행", "공장 정비지시 워크플로우를 선택하고 Run을 실행합니다.", "블록 상태 변화"],
            ["3", "우측 실행 패널", "진행률, 완료 단계, 입출력 데이터를 확인합니다.", "완료 / 오류 0"],
            ["4", "공장 게시판", "댓글 또는 정비 지시서가 등록되었는지 확인합니다.", "외부 결과 등록"],
            ["5", "관계 탐색", "요청, 설비, 고장, 정비지시 관계가 저장되었는지 보여줍니다.", "관계 그래프 표시"],
            ["6", "통합 질의", "저장된 관계와 문서 근거를 함께 질문합니다.", "AI/RAG 답변"],
            ["7", "감사 로그/DLQ", "운영 환경에서 추적과 복구가 가능함을 설명합니다.", "이력 확인"],
        ],
        widths=[0.5, 1.4, 3.4, 1.2],
    )
    add_callout(
        doc,
        "설명 멘트",
        "지금 보는 화면은 AI가 답변만 생성하는 화면이 아닙니다. 업무 요청을 받아 워크플로우로 처리하고, 결과를 외부 시스템과 온톨로지에 동시에 남기는 실행형 AI 화면입니다.",
        fill=PALETTE["amber_soft"],
    )


def scenarios(doc: Document) -> None:
    add_heading(doc, "3. 업무 시나리오 설명", 1)
    add_heading(doc, "3.1 고객사 문의 자동댓글", 2)
    add_para(doc, "고객사 게시판에 새 문의가 등록되면 온톨로지 플랫폼이 문의를 해석하고 답변 초안을 생성한 뒤 customer_mcp를 통해 댓글을 등록합니다.")
    add_table(
        doc,
        ["단계", "내용"],
        [
            ["문의 등록", "고객사 모의 게시판에 open 문의가 생성됩니다."],
            ["워크플로우 실행", "Webhook 또는 Batch 방식으로 처리합니다."],
            ["답변 생성", "RAG와 온톨로지 근거를 활용합니다."],
            ["댓글 등록", "customer_mcp comment.create를 통해 게시판에 등록합니다."],
            ["추적", "Workflow Trace와 Ontology Explorer에서 결과를 확인합니다."],
        ],
        widths=[1.4, 5.1],
    )
    doc.add_page_break()
    add_heading(doc, "3.2 공장 자동화 정비지시", 2)
    add_para(doc, "공장 자동화 시나리오는 현장 요청, 설비, 생산 라인, 공정 단계, 고장 이력, 품질 이슈를 연결해 반복 고장과 정비 필요성을 판단합니다.")
    add_table(
        doc,
        ["업무 단계", "설명"],
        [
            ["현장 요청 입력", "공장 게시판의 고장/품질 요청을 가져옵니다."],
            ["고장/품질 분류", "설비 고장인지 품질 이슈인지 판단합니다."],
            ["자산 매핑", "공장, 라인, 공정, 설비와 연결합니다."],
            ["반복 여부 확인", "동일 설비의 반복 고장 여부를 확인합니다."],
            ["정비지시 생성", "정비가 필요한 경우 작업 지시를 생성합니다."],
            ["온톨로지 저장", "요청, 설비, 고장, 정비지시 관계를 남깁니다."],
        ],
        widths=[1.8, 4.7],
    )
    add_callout(
        doc,
        "공장 시나리오 예시",
        "세종 배터리팩 공장 3번 조립 라인의 검사 카메라 촬영 품질이 10시와 11시에 반복적으로 저하되었습니다. 시스템은 동일 설비의 반복 고장인지 확인하고 정비 확인 지시 또는 품질 점검 안내를 생성합니다.",
        fill=PALETTE["green_soft"],
    )


def architecture(doc: Document) -> None:
    add_heading(doc, "4. 온톨로지, 워크플로우, RAG의 관계", 1)
    add_table(
        doc,
        ["구분", "쉬운 설명", "시스템 역할"],
        [
            ["워크플로우", "일을 처리하는 절차", "입력, 판단, 조회, 실행, 기록 단계를 수행합니다."],
            ["RAG", "문서와 지식을 찾아 답변에 활용", "정책, 매뉴얼, 지식 문서에서 근거를 검색합니다."],
            ["온톨로지", "업무 객체와 관계 지도", "요청, 설비, 답변, 정비지시의 연결을 저장합니다."],
        ],
        widths=[1.2, 2.0, 3.3],
    )
    add_callout(
        doc,
        "DB와 다른 점",
        "일반 DB는 데이터를 표로 저장하는 데 강합니다. 온톨로지는 현장 요청이 어떤 설비 고장을 보고했고, 그 고장이 어떤 생산 라인에 영향을 주었으며, 어떤 정비지시를 만들었는지를 관계로 남깁니다.",
    )
    add_heading(doc, "4.1 대표 온톨로지 관계", 2)
    add_table(
        doc,
        ["관계", "의미"],
        [
            ["ServiceRequest -> reports -> FaultEvent", "현장 요청이 고장 이벤트를 보고"],
            ["FaultEvent -> affects -> Equipment", "고장이 특정 설비에 영향"],
            ["FaultEvent -> creates -> MaintenanceTask", "고장이 정비지시를 생성"],
            ["AutoReply -> posted_as -> ExternalComment", "자동 답변이 외부 댓글로 등록"],
            ["WorkflowExecution -> generated -> AutoReply", "워크플로우 실행이 답변을 생성"],
        ],
        widths=[3.0, 3.5],
    )


def menu_overview(doc: Document) -> None:
    add_heading(doc, "5. 좌측 메뉴 전체 설명", 1)
    add_para(doc, "좌측 메뉴는 4개 그룹, 총 20개 메뉴로 구성됩니다.")
    add_table(
        doc,
        ["그룹", "메뉴", "설명"],
        [
            ["워크플로우", "워크플로우 홈", "템플릿과 최근 실행 현황을 확인합니다."],
            ["워크플로우", "템플릿 갤러리", "표준 업무 흐름을 복제해 시작합니다."],
            ["워크플로우", "빌더와 실행", "캔버스에서 업무 흐름을 편집하고 실행합니다."],
            ["워크플로우", "스킬 관리", "워크플로우에서 재사용할 실행 기능을 관리합니다."],
            ["워크플로우", "실행 추적", "실행 결과와 온톨로지 흐름을 확인합니다."],
            ["워크플로우", "승인 워크플로우", "상태 전이와 담당자 액션을 처리합니다."],
            ["워크플로우", "Writeback DLQ", "외부 등록 실패 큐를 모니터링합니다."],
            ["온톨로지", "관계 탐색", "객체와 관계를 그래프로 조회합니다."],
            ["온톨로지", "스키마 관리", "엔티티와 관계 타입을 정의합니다."],
            ["온톨로지", "인스턴스 편집", "실제 객체 데이터를 생성하고 수정합니다."],
            ["온톨로지", "관계 그래프 편집", "객체 연결과 배치를 편집합니다."],
            ["온톨로지", "RDF 워크벤치", "RDF 가져오기와 연결 데이터를 관리합니다."],
            ["질의와 분석", "대시보드", "주요 지표와 작업 큐를 확인합니다."],
            ["질의와 분석", "객체 탐색", "온톨로지 객체를 빠르게 검색합니다."],
            ["질의와 분석", "온톨로지 질의", "AI 기반 업무 질의를 수행합니다."],
            ["질의와 분석", "통합 질의", "온톨로지와 RAG를 결합해 분석합니다."],
            ["질의와 분석", "문서 RAG 질의", "문서 근거 기반 답변을 생성합니다."],
            ["질의와 분석", "SPARQL 콘솔", "정형 그래프 쿼리를 실행합니다."],
            ["운영과 검증", "감사 로그", "사용자 작업과 외부 호출 기록을 확인합니다."],
            ["운영과 검증", "통합 테스트", "질의와 근거 품질을 검증합니다."],
        ],
        widths=[1.1, 1.7, 3.7],
    )


def screen_details(doc: Document) -> None:
    add_heading(doc, "6. 핵심 화면 설명", 1)
    add_table(
        doc,
        ["화면", "사용자가 봐야 할 것", "향후 추가 기능"],
        [
            ["빌더와 실행", "실행 모드, 처리 건수, 블록 상태, 입출력, 온톨로지 매핑", "자동 포커싱, 연결선 조건 편집, 실행 전 유효성 검사, 블록별 재실행"],
            ["실행 추적", "워크플로우 실행 이력, 실패 노드, 외부 호출 결과", "실행 이력 비교, 재실행, 로그 다운로드"],
            ["관계 탐색", "요청, 답변, 설비, 고장, 정비지시 관계", "경로 하이라이트, 실행 ID 필터, 타입별 범례"],
            ["스킬 관리", "스킬 ID, 유형, 버전, 입력/출력 스키마", "스킬 테스트 실행, 버전 관리, 사용 워크플로우 목록"],
            ["감사 로그", "누가 무엇을 실행했고 어떤 외부 호출이 있었는지", "payload 마스킹, 감사 리포트, 이상 행위 탐지"],
            ["통합 테스트", "주요 시나리오 정상 동작 여부", "회귀 테스트 자동 실행, 기대/실제 결과 비교"],
        ],
        widths=[1.4, 2.7, 2.4],
    )
    add_callout(
        doc,
        "제품화 우선순위",
        "1순위는 빌더와 실행 화면의 완성도입니다. 블록 상태 시각화, 입출력 확인, 선/블록 편집, 실행 전 검증이 납품용 UX의 핵심입니다.",
        fill=PALETTE["amber_soft"],
    )


def skill_section(doc: Document) -> None:
    add_heading(doc, "7. 스킬 관리와 워크플로우 연결", 1)
    add_para(doc, "스킬은 워크플로우 노드가 실행할 수 있는 기능 단위입니다. 댓글 등록, 정비지시 생성, 온톨로지 저장, RAG 조회 같은 기능을 재사용 가능한 형태로 관리합니다.")
    add_table(
        doc,
        ["스킬 예", "역할"],
        [
            ["customer-comment-create", "고객사 게시판 댓글 등록"],
            ["factory-comment-create", "공장 게시판 댓글 등록"],
            ["factory-maintenance-create", "정비 지시 생성"],
            ["ontology-write", "온톨로지 객체/관계 저장"],
            ["rag-ontology-lookup", "RAG와 온톨로지 기반 지식 조회"],
            ["fault-recurrence-check", "반복 고장 여부 확인"],
            ["request-classify", "요청 유형 분류"],
        ],
        widths=[2.3, 4.2],
    )
    add_bullets(
        doc,
        [
            "워크플로우는 절차이고, 스킬은 그 절차 안에서 실행되는 기능입니다.",
            "스킬을 분리하면 외부 시스템 호출을 표준화하고 고객사별 구현을 교체하기 쉽습니다.",
            "향후에는 스킬 테스트 실행, 버전 관리, 권한 관리, 실행 감사 로그가 필요합니다.",
        ],
    )


def operations(doc: Document) -> None:
    add_heading(doc, "8. 데모 운영과 문제 해결", 1)
    add_table(
        doc,
        ["점검 항목", "기준"],
        [
            ["backend 8001", "ont_platform API 정상 응답"],
            ["frontend 3002", "콘솔 화면 접속 가능"],
            ["고객사 게시판 8090", "문의 등록 및 댓글 확인 가능"],
            ["공장 게시판 8091", "현장 요청 및 댓글/정비지시 확인 가능"],
            ["MCP 서버", "댓글/정비지시 등록 중계 가능"],
            ["실행 옵션", "post/dry_run, 처리 건수, 재수행 옵션 확인"],
        ],
        widths=[2.0, 4.5],
    )
    add_heading(doc, "8.1 자주 발생하는 문제", 2)
    add_table(
        doc,
        ["증상", "확인 순서"],
        [
            ["댓글이 안 달림", "워크플로우 completed 여부 -> 실행 모드 post -> MCP 실행 -> 게시판 target URL -> Workflow Trace"],
            ["계속 running", "마지막 노드 완료 여부 -> 외부 호출 응답 -> 새로고침 -> backend 로그"],
            ["진행률이 100%가 아님", "실행 대상 노드 수와 전체 노드 수 계산 차이 확인"],
            ["한글 깨짐", "브라우저 문제인지 PowerShell 출력 인코딩 문제인지 구분"],
            ["포트 충돌", "이미 실행 중인 서비스가 있는지 확인 후 기존 서버 사용 또는 재기동"],
        ],
        widths=[1.7, 4.8],
    )


def value_and_close(doc: Document) -> None:
    add_heading(doc, "9. 외부 설명용 가치 포인트", 1)
    add_table(
        doc,
        ["관점", "설명"],
        [
            ["AI 챗봇과의 차이", "답변 생성에서 끝나지 않고 실제 업무 시스템에 결과를 등록합니다."],
            ["기업 업무 적합성", "워크플로우, 권한, 감사 로그, 실패 큐 같은 운영 요소를 포함합니다."],
            ["온톨로지 가치", "업무 객체와 관계가 남아 반복 분석과 원인 추적이 가능합니다."],
            ["확장성", "스킬과 MCP 구조로 고객사별 외부 시스템 연동을 분리할 수 있습니다."],
            ["투자/PoC 가치", "기술 설명이 아니라 작동하는 업무 자동화 PoC로 보여줄 수 있습니다."],
        ],
        widths=[1.7, 4.8],
    )
    add_callout(
        doc,
        "마무리 멘트",
        "이 솔루션은 온톨로지, 워크플로우, RAG를 따로 보여주는 기술 데모가 아니라, 현장 요청을 처리하고 외부 시스템에 결과를 남기며 그 관계를 다시 추적할 수 있는 업무 실행형 AI 플랫폼의 원형입니다.",
        fill=PALETTE["green_soft"],
    )


def add_source_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "부록. 원본 매뉴얼 출처", 1)
    add_table(
        doc,
        ["문서", "역할"],
        [
            ["10_시나리오_통합_실행_매뉴얼.md", "전체 시나리오 실행 공통 절차"],
            ["11_고객사_문의_자동댓글_시나리오_매뉴얼.md", "고객사 문의 자동댓글"],
            ["12_공장자동화_정비지시_시나리오_매뉴얼.md", "공장 자동화 정비지시"],
            ["13_온톨로지_워크플로우_RAG_연계_매뉴얼.md", "온톨로지/워크플로우/RAG 관계"],
            ["14_스킬_관리_및_워크플로우_연결_매뉴얼.md", "스킬 관리와 워크플로우 연결"],
            ["15_데모_운영_점검_트러블슈팅_매뉴얼.md", "데모 운영과 문제 해결"],
            ["16_좌측_메뉴_전체_설명_매뉴얼.md", "좌측 메뉴 전체 설명"],
            ["17_메뉴별_화면_상세_및_향후기능_매뉴얼.md", "메뉴별 화면 상세와 향후 기능"],
        ],
        widths=[3.6, 2.9],
    )


def build() -> None:
    doc = Document()
    configure_doc(doc)
    cover(doc)
    executive_summary(doc)
    demo_flow(doc)
    scenarios(doc)
    architecture(doc)
    menu_overview(doc)
    screen_details(doc)
    skill_section(doc)
    operations(doc)
    value_and_close(doc)
    add_source_appendix(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
