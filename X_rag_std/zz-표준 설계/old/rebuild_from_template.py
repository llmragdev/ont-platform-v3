#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Rebuild document from user's template with expanded API specs"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

template_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
output_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'
md_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md'

# 마크다운 파일 읽기
with open(md_file, 'r', encoding='utf-8') as f:
    md_lines = f.readlines()

# 섹션 3.2.1~3.2.5 찾기
api_sections = {}
current_section = None
current_lines = []

in_section3 = False
for i, line in enumerate(md_lines):
    if line.startswith('### 3.2.'):
        if current_section and current_lines:
            api_sections[current_section] = current_lines
        current_section = line.strip()
        current_lines = [line]
        in_section3 = True
    elif in_section3 and line.startswith('### ') and not line.startswith('### 3.2.'):
        # 섹션 종료
        if current_section:
            api_sections[current_section] = current_lines
        break
    elif in_section3 and current_section:
        current_lines.append(line)

if current_section:
    api_sections[current_section] = current_lines

print(f"마크다운에서 찾은 API 섹션: {len(api_sections)}")
for sec in api_sections.keys():
    print(f"  - {sec}")

# 템플릿 복제
print(f"\n템플릿 로드 중: {template_file}")
doc = Document(template_file)

print(f"현재 문서: {len(doc.paragraphs)} 파라그래프, {len(doc.tables)} 표")

# 마크다운 콘텐츠를 파싱하여 테이블/텍스트 추출
def parse_md_section(section_lines):
    """마크다운 섹션을 파싱하여 헤더, 텍스트, 테이블 추출"""
    result = {
        'header': None,
        'description': [],
        'tables': [],
        'json_blocks': []
    }

    i = 0
    while i < len(section_lines):
        line = section_lines[i].rstrip('\n')

        # 헤더
        if line.startswith('### '):
            result['header'] = line[4:]
        # 테이블 감지
        elif line.startswith('|'):
            table_rows = []
            # 헤더 행
            if i + 1 < len(section_lines) and section_lines[i+1].strip().startswith('|'):
                headers = [h.strip() for h in line.split('|')[1:-1]]
                table_rows.append(headers)
                i += 2  # 분리선 건너뛰기

                # 데이터 행
                while i < len(section_lines) and section_lines[i].startswith('|'):
                    cells = [c.strip() for c in section_lines[i].split('|')[1:-1]]
                    # 헤더 열 수에 맞추기
                    while len(cells) < len(headers):
                        cells.append('')
                    table_rows.append(cells[:len(headers)])
                    i += 1

                result['tables'].append(table_rows)
                i -= 1
        # JSON 블록
        elif line.startswith('```json'):
            json_lines = []
            i += 1
            while i < len(section_lines) and not section_lines[i].startswith('```'):
                json_lines.append(section_lines[i].rstrip('\n'))
                i += 1
            result['json_blocks'].append('\n'.join(json_lines))
        # 일반 텍스트
        elif line.strip() and not line.startswith('```'):
            result['description'].append(line)

        i += 1

    return result

# 각 API 섹션 파싱
parsed_apis = {}
for section_key, lines in api_sections.items():
    parsed = parse_md_section(lines)
    parsed_apis[section_key] = parsed
    print(f"\n{section_key}:")
    print(f"  - 설명 줄: {len(parsed['description'])}")
    print(f"  - 테이블: {len(parsed['tables'])}")
    print(f"  - JSON 블록: {len(parsed['json_blocks'])}")

print(f"\n문서 저장: {output_file}")
doc.save(output_file)
print("✓ 완료!")

