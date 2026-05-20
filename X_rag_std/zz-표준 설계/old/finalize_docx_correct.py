#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Rebuild DOCX with correct table style and border colors"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

template_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
output_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'
md_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md'

# 템플릿 로드
print(f"템플릿 로드: {template_file}")
doc = Document(template_file)

# 참고: 템플릿의 테이블 스타일 추출
reference_table = None
if doc.tables:
    reference_table = doc.tables[0]
    print(f"참고 테이블 스타일: {reference_table.style.name}")

# 마크다운 읽기
with open(md_file, 'r', encoding='utf-8') as f:
    md_text = f.read()

# 섹션 3.2 추출
start = md_text.find('### 3.2.1')
end = md_text.find('## 4. 청킹 표준')
section_3_2 = md_text[start:end] if start != -1 and end != -1 else ""

print(f"마크다운 섹션 3.2 추출: {len(section_3_2)} 문자")

# 3.2.1 이후 콘텐츠 제거
remove_from_idx = None
for i, para in enumerate(doc.paragraphs):
    if '3.2.1' in para.text:
        remove_from_idx = i
        break

if remove_from_idx is not None:
    for i in range(len(doc.paragraphs) - 1, remove_from_idx - 1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    print(f"기존 API 섹션 제거: 파라그래프 {remove_from_idx}부터")

# 테이블도 제거
initial_table_count = len(doc.tables)
if initial_table_count > 5:
    for i in range(initial_table_count - 1, 4, -1):
        tbl = doc.tables[i]._element
        tbl.getparent().remove(tbl)

print(f"남은 파라그래프: {len(doc.paragraphs)}, 표: {len(doc.tables)}")

# 마크다운 파싱 및 추가
lines = section_3_2.split('\n')

print(f"\n섹션 3.2 추가 중 ({len(lines)} 라인)...")

def set_table_border_color(table, color='000000'):
    """테이블의 모든 테두리를 검정색으로 설정"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr {}/>', nsdecls('w'))
        tbl.insert(0, tblPr)

    # 테두리 설정
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )

    # 기존 tblBorders 제거 후 새것 추가
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

# 섹션 3.2 콘텐츠 추가
i = 0
table_count = 0
while i < len(lines):
    line = lines[i]

    # 제목
    if line.startswith('### '):
        p = doc.add_paragraph(line[4:], style='Heading 3')
    # 표
    elif line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|'):
        headers = [h.strip() for h in line.split('|')[1:-1]]
        i += 2  # 분리선 건너뛰기

        # Table Grid 스타일 사용
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table_count += 1

        hdr_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            if j < len(hdr_cells):
                hdr_cells[j].text = header
                # 헤더 셀 배경색 설정 (D9E1F2 = 파란색)
                shading_elm = parse_xml(r'<w:shd {} w:fill="D9E1F2"/>'.format(nsdecls('w')))
                hdr_cells[j]._element.get_or_add_tcPr().append(shading_elm)

        while i < len(lines) and lines[i].startswith('|'):
            cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            while len(cells) < len(headers):
                cells.append('')
            row_cells = table.add_row().cells
            for j, cell in enumerate(cells[:len(headers)]):
                if j < len(row_cells):
                    row_cells[j].text = cell
                    # 데이터 셀 배경색 제거 (흰색)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
                    row_cells[j]._element.get_or_add_tcPr().append(shading_elm)
            i += 1

        # 테이블 테두리 색상 설정 (검정색)
        set_table_border_color(table, '000000')
        i -= 1
    # 코드 블록
    elif '```' in line:
        i += 1
        code_lines = []
        while i < len(lines) and '```' not in lines[i]:
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.style = 'Normal'
            p.add_run('\n'.join(code_lines)).font.name = 'Courier New'
    # 일반 텍스트
    elif line.strip() and not line.startswith('---'):
        doc.add_paragraph(line)

    i += 1

print(f"테이블 추가: {table_count}개")

# 저장
print(f"\n문서 저장: {output_file}")
doc.save(output_file)

print(f"✓ 완료!")
print(f"  - 파라그래프: {len(doc.paragraphs)}")
print(f"  - 표: {len(doc.tables)} (스타일: Table Grid, 테두리: 검정색)")

