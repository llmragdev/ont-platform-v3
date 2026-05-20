#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Direct fix - manipulate paragraph text directly"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

fixes_done = []

print("=" * 80)
print("직접 수정 (Paragraph 전체 재작성)")
print("=" * 80)

# 모든 파라그래프 순회하며 직접 수정
for i, para in enumerate(doc.paragraphs):
    old_text = para.text
    new_text = old_text

    # 1. 백터DB → 벡터DB
    if '백터DB' in new_text:
        new_text = new_text.replace('백터DB', '벡터DB')
        if new_text != old_text:
            fixes_done.append({
                'line': i,
                'type': '오타',
                'from': '백터DB',
                'to': '벡터DB'
            })
            print(f"  ✓ 라인 {i}: 백터DB → 벡터DB")

    # 2. 라인 147 근처에서 DELETE 관련 문구 찾기
    if i == 147 or (140 <= i <= 150 and 'DELETE' in new_text):
        if '3.2.4 문서재업로드' in new_text and 'DELETE' in new_text:
            new_text = new_text.replace('3.2.4 문서재업로드', '3.2.5 문서삭제')
            if new_text != old_text:
                fixes_done.append({
                    'line': i,
                    'type': '중복섹션',
                    'from': old_text[:70],
                    'to': new_text[:70]
                })
                print(f"  ✓ 라인 {i}: 문서재업로드 → 문서삭제")

    # 텍스트가 변경되었으면 파라그래프 재작성
    if new_text != old_text:
        # 스타일 유지하면서 내용만 변경
        style_backup = para.style
        para.clear()
        para.style = style_backup
        para.add_run(new_text)

# 저장
print(f"\n[최종] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료: {guide_file}")

# 요약
print("\n" + "=" * 80)
print("직접 수정 결과")
print("=" * 80)

if fixes_done:
    print(f"\n총 {len(fixes_done)}개 추가 수정:\n")
    for idx, fix in enumerate(fixes_done, 1):
        print(f"[{idx}] {fix['type']}")
        print(f"    라인: {fix['line']}")
        print(f"    변경: {fix['from']} → {fix['to']}")
else:
    print("\n추가 수정 사항 없음 (이미 수정됨)")

print("\n" + "=" * 80)

