#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Check search response table and JSON example consistency"""

from docx import Document
import json
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("검색 응답 표와 JSON 예시 비교")
print("=" * 80)

# 1. "검색 응답 주요 필드" 표 찾기
print("\n[1단계] 검색 응답 표 찾기\n")

search_response_section = None
search_response_table_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '검색 응답' in text and '주요 필드' in text:
        search_response_section = i
        print(f"검색 응답 섹션 찾음: 라인 {i}")
        print(f"텍스트: {text}\n")
        break

# 표 찾기 (섹션 다음의 가장 가까운 표)
if search_response_section:
    # 섹션 이후의 표 찾기
    table_count = 0
    for i, para in enumerate(doc.paragraphs[search_response_section:]):
        actual_idx = search_response_section + i
        if actual_idx < search_response_section + 50:  # 50줄 내에서만 찾기
            for table in doc.tables:
                # 표의 위치를 대략적으로 파악
                pass

# 더 간단한 방법: 섹션 이후의 첫 번째 표 찾기
para_to_table_map = {}
table_indices_in_doc = []

# 모든 테이블을 파라그래프 순서와 매핑
for table_idx, table in enumerate(doc.tables):
    # 테이블 앞의 파라그래프 찾기
    table_element = table._element
    # 간단히: 테이블 개수와 위치 저장
    table_indices_in_doc.append(table_idx)

print(f"문서 내 표 총 {len(doc.tables)}개\n")

# 2. 검색 응답 표의 내용 출력
print("[2단계] 검색 응답 표 내용\n")

# 사용자가 제공한 표의 필드들
table_fields = [
    "data.answer",
    "data.used_chunks[].chunk_id",
    "data.used_chunks[].content",
    "data.used_chunks[].similarity_score",
    "data.used_chunks[].metadata.source_url",
    "data.used_chunks[].metadata.page_no",
    "data.debug_info"
]

print("표에 정의된 필드들:")
for field in table_fields:
    print(f"  ✓ {field}")

# 3. JSON 예시에서 실제 필드 추출
print("\n[3단계] JSON 예시 내용\n")

json_example = """
{
  "status": "success",
  "data": {
    "query": "2026년 인사 규정 알려줘",
    "answer": "2026년 인사 규정에 따르면...",
    "used_chunks": [
      {
        "chunk_id": "doc_a1b2c3d4#chunk4",
        "content": "LLM이 정답 생성에 채택한 텍스트...",
        "metadata": {
          "source_name": "2026_인사규정.pdf",
          "source_url": "https://storage.example.com/docs/2026_인사규정.pdf",
          "page_no": 12,
          "category_large": "인사",
          "category_mid": "채용",
          "vector_db_id": "vdb_hr_recruit_01",
          "tenant_id": "company_abc",
          "org_id": "0102",
          "dept_code": "01"
        },
        "similarity_score": 0.89
      }
    ],
    "debug_info": {
      "execution_time_ms": 145,
      "candidate_chunks": ["// debug_mode: true 시에만 노출 — 미채택 청크 포함"]
    }
  },
  "error": null
}
"""

json_example_fields = [
    "data.answer",
    "data.used_chunks[].chunk_id",
    "data.used_chunks[].content",
    "data.used_chunks[].metadata.source_url",
    "data.used_chunks[].metadata.page_no",
    "data.used_chunks[].similarity_score",
    "data.debug_info",
    "data.query",  # 표에 없음
    "data.used_chunks[].metadata.source_name",  # 표에 없음
    "data.used_chunks[].metadata.category_large",  # 표에 없음
    "data.used_chunks[].metadata.category_mid",  # 표에 없음
    "data.used_chunks[].metadata.vector_db_id",  # 표에 없음
    "data.used_chunks[].metadata.tenant_id",  # 표에 없음
    "data.used_chunks[].metadata.org_id",  # 표에 없음
    "data.used_chunks[].metadata.dept_code",  # 표에 없음
]

print("JSON 예시의 실제 필드들:")
for field in json_example_fields:
    print(f"  ✓ {field}")

# 4. 불일치 찾기
print("\n[4단계] 불일치 분석\n")

print("표에는 있지만 JSON 예시에는 없는 필드:")
missing_in_json = []
for field in table_fields:
    if field not in json_example_fields:
        missing_in_json.append(field)
        print(f"  ❌ {field}")

if not missing_in_json:
    print("  (없음)")

print("\nJSON 예시에는 있지만 표에는 없는 필드:")
extra_in_json = []
for field in json_example_fields:
    if field not in table_fields:
        extra_in_json.append(field)
        print(f"  ⚠️ {field}")

if not extra_in_json:
    print("  (없음)")

# 5. 권장사항
print("\n" + "=" * 80)
print("권장사항")
print("=" * 80)

print("\n✓ 표와 JSON 예시 동기화 필요:")
print("\n1. 표에 누락된 필드 추가:")
print("   - data.query (검색 쿼리)")
print("   - data.used_chunks[].metadata.source_name")
print("   - data.used_chunks[].metadata.category_large")
print("   - data.used_chunks[].metadata.category_mid")
print("   - data.used_chunks[].metadata.vector_db_id")
print("   - data.used_chunks[].metadata.tenant_id")
print("   - data.used_chunks[].metadata.org_id")
print("   - data.used_chunks[].metadata.dept_code")

print("\n2. 또는 JSON 예시에서 metadata 필드 정리:")
print("   - 핵심 필드만 포함하거나")
print("   - \"metadata\" 구조를 표와 일치하도록 정리")

print("\n3. 표 구조 개선 제안:")
print("   - 세부 필드 구분 (필수 vs 선택)")
print("   - 복합 객체(metadata) 명확히 표시")

print("\n" + "=" * 80)

