#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Analyze template in detail - tables, colors, formatting"""

from docx import Document
from docx.oxml.ns import qn
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

template_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
doc = Document(template_file)

print("=" * 80)
print("템플릿 상세 분석 (_내가편집.docx)")
print("=" * 80)

# 각 API 섹션 분석
print("\n각 API 섹션 구조 분석:")
api_sections = []
current_api = None

for i, para in enumerate(doc.paragraphs):
    if '3.2.' in para.text and para.style.name == 'Heading 3':
        current_api = {
            'heading': para.text,
            'heading_idx': i,
            'description': [],
            'table_indices': []
        }
        api_sections.append(current_api)
    elif current_api and i > current_api['heading_idx']:
        # 다음 섹션까지 설명 수집
        if para.style.name == 'Heading 3':
            break
        if para.text.strip() and para.style.name in ['Normal', 'List Bullet']:
            current_api['description'].append(para.text)

# 테이블과 매핑
for api in api_sections:
    print(f"\n{api['heading']}:")
    print(f"  위치: 파라그래프 {api['heading_idx']}")
    print(f"  설명:")
    for desc in api['description'][:3]:  # 처음 3개만
        print(f"    - {desc[:70]}")

# 테이블 포맷 분석
print("\n" + "=" * 80)
print("테이블 포맷 분석")
print("=" * 80)

for t_idx, table in enumerate(doc.tables[:3]):
    print(f"\n표 {t_idx}:")
    print(f"  크기: {len(table.rows)}행 x {len(table.columns)}열")
    print(f"  스타일: {table.style.name if table.style else 'None'}")

    # 테이블 테두리 색상 확인
    try:
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is not None:
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                print(f"  테두리 설정: 있음")
                for border_name in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                    border_elem = tblBorders.find(qn(f'w:{border_name}'))
                    if border_elem is not None:
                        color = border_elem.get(qn('w:color'))
                        size = border_elem.get(qn('w:sz'))
                        val = border_elem.get(qn('w:val'))
                        print(f"    {border_name}: color={color}, size={size}, val={val}")
    except Exception as e:
        print(f"  테두리 분석 오류: {e}")

    # 헤더 행 배경색 확인
    if len(table.rows) > 0:
        header_cell = table.rows[0].cells[0]
        try:
            tcPr = header_cell._element.tcPr
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    print(f"  헤더 셀 배경색: {fill}")
        except:
            pass

print("\n완료!")
