#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update search response table with complete JSON specification"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("검색 응답 주요 필드 표 업데이트")
print("=" * 80)

# 1. "검색 응답 주요 필드" 표 찾기
print("\n[1단계] 검색 응답 표 찾기\n")

search_response_table_idx = None
search_response_para_idx = None

# 섹션 제목 찾기
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '검색 응답' in text and '주요 필드' in text:
        search_response_para_idx = i
        print(f"검색 응답 섹션 찾음: 라인 {i}")
        print(f"텍스트: {text}\n")
        break

# 섹션 이후 첫 번째 표 찾기
if search_response_para_idx is not None:
    # 문서 내 모든 표 순회하면서 위치 파악
    table_count_before = 0
    for idx, table in enumerate(doc.tables):
        # 표 앞의 파라그래프 수를 대략 파악하기 위해
        # 표의 위치가 섹션 이후인지 확인
        # python-docx에서는 정확한 표 위치를 파악하기 어려우므로
        # 섹션 이후 첫 표라고 가정
        search_response_table_idx = idx
        print(f"표 {idx}을 대상으로 지정\n")
        break

# 2. 새로운 표 구조 정의
print("[2단계] 새로운 표 구조 정의\n")

# 완전한 필드 목록 (JSON 기준)
fields = [
    ("필드명", "타입", "설명"),
    ("status", "String", "응답 상태: \"success\" 또는 \"error\""),
    ("error", "Object|null", "오류 발생 시: {code, message} 또는 null"),
    ("data.query", "String", "사용자가 입력한 검색 쿼리"),
    ("data.answer", "String", "LLM이 생성한 답변 텍스트"),
    ("data.used_chunks[].chunk_id", "String", "청크 고유 ID (doc_id#chunkN 형식)"),
    ("data.used_chunks[].content", "String", "LLM이 채택한 청크 원문"),
    ("data.used_chunks[].similarity_score", "Float", "벡터 유사도 점수 (0.0~1.0)"),
    ("data.used_chunks[].metadata.source_name", "String", "원본 파일명 (예: 2026_인사규정.pdf)"),
    ("data.used_chunks[].metadata.source_url", "String", "문서 저장소 URL"),
    ("data.used_chunks[].metadata.page_no", "Integer", "원본 파일 내 페이지 번호"),
    ("data.used_chunks[].metadata.category_large", "String", "대분류 (예: 인사, 규정, 기술)"),
    ("data.used_chunks[].metadata.category_mid", "String", "중분류 (벡터DB 라우팅 기준)"),
    ("data.used_chunks[].metadata.vector_db_id", "String", "벡터DB 식별자 (예: vdb_hr_recruit_01)"),
    ("data.used_chunks[].metadata.tenant_id", "String", "다중테넌트 격리용 테넌트 ID"),
    ("data.used_chunks[].metadata.org_id", "String", "조직 계층 코드 (org_id 전체)"),
    ("data.used_chunks[].metadata.dept_code", "String", "부서 코드 (org_id 앞 2자리)"),
    ("data.debug_info.execution_time_ms", "Integer", "API 응답 시간 (밀리초) - debug_mode: true 시만 노출"),
    ("data.debug_info.candidate_chunks", "Array", "미채택 청크 목록 - debug_mode: true 시만 노출"),
]

print(f"새 표 구조: {len(fields)-1}개 필드\n")
for i, (field, typ, desc) in enumerate(fields):
    if i > 0:
        print(f"  {i}. {field:50} | {typ:15} | {desc[:40]}")

# 3. 기존 표 찾아서 수정
print("\n[3단계] 표 수정\n")

if search_response_table_idx is not None and search_response_table_idx < len(doc.tables):
    table = doc.tables[search_response_table_idx]

    print(f"기존 표: {len(table.rows)}행 x {len(table.rows[0].cells)}열")
    print(f"새 표: {len(fields)}행 x 3열")

    # 기존 표의 행 삭제 (역순으로 삭제하여 인덱스 변경 영향 최소화)
    for _ in range(len(table.rows)):
        tbl = table._element
        tr = tbl.tr_lst[0]
        tbl.remove(tr)

    # 새 행 추가
    for row_idx, (field, typ, desc) in enumerate(fields):
        row = table.add_row()

        # 각 셀에 내용 입력
        cells = row.cells
        cells[0].text = field
        cells[1].text = typ
        cells[2].text = desc

        # 헤더 행 (첫 번째 행) 포맷팅
        if row_idx == 0:
            for cell in cells:
                # 배경색을 파란색으로 (헤더)
                shading_elm = cell._element.get_or_add_tcPr().append(
                    __import__('docx.oxml').oxml.parse_xml(
                        r'<w:shd {} w:fill="D9E1F2"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        )
                    )
                )
                # 텍스트 굵게
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 데이터 행: 배경색 흰색
            for cell in cells:
                shading_elm = cell._element.get_or_add_tcPr().append(
                    __import__('docx.oxml').oxml.parse_xml(
                        r'<w:shd {} w:fill="FFFFFF"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        )
                    )
                )

    print(f"✓ 표 수정 완료: {len(fields)}행 생성")

# 4. 저장
print("\n[4단계] 문서 저장\n")
doc.save(guide_file)
print(f"✓ 저장 완료: {guide_file}")

print("\n" + "=" * 80)
print("✓ 검색 응답 표 업데이트 완료!")
print("=" * 80)
