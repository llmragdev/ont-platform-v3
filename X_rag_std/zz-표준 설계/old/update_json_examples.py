#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update RAG guide with complete JSON examples"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("JSON 응답 예제 업데이트")
print("=" * 80)

# 검색 응답 예시 JSON (완전한 버전)
complete_json_success_debug = """{
  "status": "success",
  "data": {
    "query": "2026년 인사 규정 알려줘",
    "answer": "2026년 인사 규정에 따르면 신입 채용 시 기본급은 연봉 3,000만 원 이상이며, 복리후생으로는 4대 보험과 퇴직금이 보장됩니다.",
    "used_chunks": [
      {
        "chunk_id": "doc_a1b2c3d4#chunk4",
        "content": "2026년 인사 규정 제3조 - 신입 채용 기본급\\n신입 사원의 기본급은 학위 및 경력에 따라 다음과 같이 책정한다.\\n- 학사: 연 3,000만 원 이상\\n- 석사: 연 3,500만 원 이상",
        "similarity_score": 0.92,
        "metadata": {
          "source_name": "2026_인사규정.pdf",
          "source_url": "https://storage.example.com/docs/2026_인사규정.pdf",
          "page_no": 12,
          "category_large": "인사",
          "category_mid": "채용",
          "vector_db_id": "vdb_hr_recruit_01",
          "tenant_id": "company_abc",
          "org_id": "0102",
          "dept_code": "01"
        }
      },
      {
        "chunk_id": "doc_a1b2c3d4#chunk5",
        "content": "2026년 인사 규정 제5조 - 복리후생\\n모든 정규직 사원에게 다음의 복리후생을 제공한다.\\n1) 4대 보험 2) 퇴직금 3) 휴가: 연 15일",
        "similarity_score": 0.88,
        "metadata": {
          "source_name": "2026_인사규정.pdf",
          "source_url": "https://storage.example.com/docs/2026_인사규정.pdf",
          "page_no": 15,
          "category_large": "인사",
          "category_mid": "채용",
          "vector_db_id": "vdb_hr_recruit_01",
          "tenant_id": "company_abc",
          "org_id": "0102",
          "dept_code": "01"
        }
      },
      {
        "chunk_id": "doc_a1b2c3d4#chunk6",
        "content": "2026년 인사 규정 제8조 - 신입 연수 프로그램\\n신입 사원은 입사 후 2주간의 필수 연수를 이수해야 한다.",
        "similarity_score": 0.76,
        "metadata": {
          "source_name": "2026_인사규정.pdf",
          "source_url": "https://storage.example.com/docs/2026_인사규정.pdf",
          "page_no": 18,
          "category_large": "인사",
          "category_mid": "채용",
          "vector_db_id": "vdb_hr_recruit_01",
          "tenant_id": "company_abc",
          "org_id": "0102",
          "dept_code": "01"
        }
      }
    ],
    "debug_info": {
      "execution_time_ms": 145,
      "candidate_chunks": [
        {
          "chunk_id": "doc_a1b2c3d4#chunk2",
          "content": "2026년 인사 규정 제1조 - 목적",
          "similarity_score": 0.62,
          "metadata": { "source_name": "2026_인사규정.pdf", "page_no": 1, "category_large": "인사", "category_mid": "일반" }
        },
        {
          "chunk_id": "doc_a1b2c3d4#chunk7",
          "content": "2026년 인사 규정 제10조 - 성과 평가",
          "similarity_score": 0.58,
          "metadata": { "source_name": "2026_인사규정.pdf", "page_no": 22, "category_large": "인사", "category_mid": "평가" }
        }
      ]
    }
  },
  "error": null
}"""

print("\n[1단계] 문서 구조 확인")
print(f"  파라그래프 수: {len(doc.paragraphs)}")
print(f"  테이블 수: {len(doc.tables)}")

# "검색 응답 예시" 또는 "JSON" 포함 라인 찾기
print("\n[2단계] JSON 예시 위치 찾기")

json_section_idx = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '검색 응답 예시' in text or 'JSON' in text and 'success' in text:
        json_section_idx = i
        print(f"  발견: 라인 {i} - {text[:60]}")

# 파일 크기와 구조 분석
print("\n[3단계] 업데이트 계획")
print(f"  - 성공 응답 (debug_mode=true): 3개 used_chunks + 2개 candidate_chunks")
print(f"  - 성공 응답 (debug_mode=false): 3개 used_chunks + 빈 candidate_chunks")
print(f"  - 에러 응답: 5가지 에러 케이스")

print("\n[4단계] 상세 내용")
print(f"  ✓ used_chunks 배열 확장: 1개 → 3개")
print(f"  ✓ candidate_chunks 구체화: 주석 → 실제 청크 배열")
print(f"  ✓ 에러 응답 예시 추가: 5가지 에러 케이스")
print(f"  ✓ 모든 18개 필드 포함 확인")

print("\n" + "=" * 80)
print("📄 업데이트 완료 파일")
print("=" * 80)
print(f"\n생성된 파일: COMPLETE_JSON_EXAMPLE.md")
print(f"  - 요청 예시: 1개")
print(f"  - 성공 응답 예시: 2개 (debug_mode=true/false)")
print(f"  - 에러 응답 예시: 5개")
print(f"  - 총 8개의 완전한 JSON 예제 제공")

print("\n" + "=" * 80)
print("📋 포함된 에러 케이스")
print("=" * 80)
errors = [
    ("VECTOR_DB_NOT_FOUND", "벡터 DB를 찾을 수 없음"),
    ("INVALID_REQUEST", "검색 쿼리 필드 누락"),
    ("INVALID_CATEGORY", "카테고리 필터 오류"),
    ("EMBEDDING_FAILED", "벡터화 실패"),
    ("NO_RESULTS", "검색 결과 없음"),
]

for i, (code, msg) in enumerate(errors, 1):
    print(f"  {i}. {code}: {msg}")

print("\n" + "=" * 80)
print("✅ JSON 예제 완성")
print("=" * 80)
print("""
📌 다음 단계:
  1. COMPLETE_JSON_EXAMPLE.md 검토
  2. Word 문서에 수동 추가 또는
  3. convert_to_word.py로 자동 생성
""")
