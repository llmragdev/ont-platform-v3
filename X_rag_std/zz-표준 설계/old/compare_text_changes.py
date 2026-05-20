#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Compare original and modified files - text changes only"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

original_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1 - 원본.docx'
modified_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'

try:
    doc_orig = Document(original_file)
    doc_mod = Document(modified_file)
except Exception as e:
    print(f"파일 읽기 오류: {e}")
    sys.exit(1)

print("=" * 80)
print("원본 vs 수정본 - 문구 변경 비교")
print("=" * 80)

text_changes = []

# 파라그래프 텍스트만 비교 (스타일 제외)
print(f"\n원본: {len(doc_orig.paragraphs)} 파라그래프")
print(f"수정본: {len(doc_mod.paragraphs)} 파라그래프")

# 각 라인의 텍스트만 비교
for i in range(min(len(doc_orig.paragraphs), len(doc_mod.paragraphs))):
    orig_text = doc_orig.paragraphs[i].text.strip()
    mod_text = doc_mod.paragraphs[i].text.strip()

    if orig_text != mod_text and orig_text and mod_text:  # 둘 다 비어있지 않을 때만
        text_changes.append({
            'line': i,
            'from': orig_text,
            'to': mod_text
        })

# 결과 출력
print(f"\n" + "=" * 80)
print(f"문구 변경 (텍스트만): 총 {len(text_changes)}개")
print("=" * 80)

if text_changes:
    for idx, change in enumerate(text_changes, 1):
        print(f"\n[{idx}] 라인 {change['line']}")
        print(f"  변경 전: {change['from']}")
        print(f"  변경 후: {change['to']}")
else:
    print("\n문구 변경 없음")

print("\n" + "=" * 80)

