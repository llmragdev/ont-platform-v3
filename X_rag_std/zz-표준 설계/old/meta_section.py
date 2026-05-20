from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows_data):
    table = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    table.style = 'Table Grid'
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        set_cell_bg(cell, 'D9E1F2')
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.cell(ri+1, ci)
            cell.text = ''
            cell.paragraphs[0].add_run(val).font.size = Pt(9.5)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    doc.add_paragraph()

# ── 제목 ──────────────────────────────────────────────────
doc.add_heading('6. 메타데이터 관리', level=2)

# ── 6.1 메타데이터란 ──────────────────────────────────────
doc.add_heading('6.1 메타데이터란', level=3)

p = doc.add_paragraph(
    '메타데이터(Metadata)란 "데이터에 대한 데이터"로, 파일·문서·콘텐츠의 내용을 설명하는 부가 속성 정보입니다. '
    '메타데이터는 특정 기술에 국한된 개념이 아니며, 우리가 다루는 모든 파일에 이미 존재합니다.'
)

add_table(doc,
    ['파일 유형', '메타데이터 예시', '저장 위치'],
    [
        ['PDF',         '제목, 작성자, 작성일, 페이지 수, 키워드',       '파일 내 Document Properties'],
        ['Word(DOCX)',  '제목, 작성자, 수정일, 회사명',                  '파일 내 Document Properties'],
        ['HTML',        '제목(<title>), 설명·키워드(<meta> 태그)',        '페이지 헤더 태그'],
        ['이미지(JPEG)', '촬영일시, GPS 좌표, 카메라 모델',              'EXIF 데이터'],
        ['음악(MP3)',    '제목, 아티스트, 앨범, 발매연도',               'ID3 태그'],
    ]
)

p = doc.add_paragraph(
    'RAG 시스템에서도 동일한 개념이 적용됩니다. 문서를 청크(chunk) 단위로 분할하여 '
    '벡터 DB에 저장할 때, 각 청크에 부가 속성을 함께 저장합니다. '
    '벡터 DB는 이 속성을 검색 결과와 함께 자동으로 반환하며, 필터 조건으로도 활용합니다.'
)
doc.add_paragraph()

# ── 6.2 자동 생성 vs 명시 설정 ────────────────────────────
doc.add_heading('6.2 자동 생성 vs 명시 설정', level=3)

p = doc.add_paragraph(
    '청크에 부착되는 메타데이터는 생성 방식에 따라 두 가지로 구분됩니다.'
)

add_table(doc,
    ['구분', '속성', '생성 주체', '설명'],
    [
        ['자동 생성', 'doc_id',     '시스템', '업로드 시 자동 발급되는 문서 고유 ID'],
        ['자동 생성', 'chunk_id',   '파이프라인', '{doc_id}#chunk{n} 형식으로 자동 생성'],
        ['자동 생성', 'dept_code',  '파이프라인', 'org_id 앞 2자리 자동 파생 (수동 입력 금지)'],
        ['자동 생성', 'page_no',    'PDF 파서', 'PDF 파싱 시 실제 페이지 번호 자동 추출'],
        ['자동 생성', 'created_at', '시스템', '저장 시점 UTC 시간 자동 기록'],
        ['명시 설정', 'tenant_id',  '요청 헤더', '호출 측이 X-Tenant-ID 헤더로 전달'],
        ['명시 설정', 'org_id',     '요청 헤더', '호출 측이 X-Org-ID 헤더로 전달 (필수)'],
        ['명시 설정', 'category_mid/large', '업로드 폼', '업로드 시 사용자가 선택·입력'],
        ['명시 설정', 'source_url', '파이프라인', '저장 완료 후 파이프라인이 URL 생성·설정'],
        ['명시 설정', 'vector_db_id', '매핑 규칙', 'category_mid → 벡터DB 매핑으로 자동 결정'],
    ]
)

# ── 6.3 처리 흐름 ─────────────────────────────────────────
doc.add_heading('6.3 메타데이터 처리 흐름', level=3)

p = doc.add_paragraph('문서 업로드부터 검색까지 메타데이터가 어떻게 흐르는지 코드로 확인합니다.')
doc.add_paragraph()

p = doc.add_paragraph('① 업로드 시 — 메타데이터 구성 (개발자 명시 설정 + 자동 파생)')
p.runs[0].bold = True

add_code(doc,
'''def build_chunk_metadata(tenant_id, org_id, category_mid, doc_id, chunk_idx, page_no):
    return {
        # 명시 설정 (호출 측 전달)
        "tenant_id":      tenant_id,          # "company_abc"
        "org_id":         org_id,             # "0102"
        # 자동 파생
        "dept_code":      org_id[:2],         # "01"
        # 명시 설정 (업로드 폼)
        "category_mid":   category_mid,       # "채용"
        # 자동 생성
        "doc_id":         doc_id,             # "doc_a1b2c3d4"
        "chunk_id":       f"{doc_id}#chunk{chunk_idx}",
        "page_no":        page_no,            # PDF 파서 추출값
    }''')

p = doc.add_paragraph('② 벡터 DB 저장 — metadatas 파라미터로 전달')
p.runs[0].bold = True

add_code(doc,
'''collection.add(
    ids        = ["doc_a1b2c3d4#chunk0", "doc_a1b2c3d4#chunk1"],
    embeddings = [[0.12, 0.34, ...], [0.56, 0.78, ...]],
    documents  = ["청크 원문 텍스트1", "청크 원문 텍스트2"],
    metadatas  = [
        {"tenant_id": "company_abc", "org_id": "0102", "dept_code": "01",
         "category_mid": "채용", "page_no": 3, ...},
        {"tenant_id": "company_abc", "org_id": "0102", "dept_code": "01",
         "category_mid": "채용", "page_no": 5, ...},
    ]
)''')

p = doc.add_paragraph('③ 검색 시 — 메타데이터 필터 조건으로 활용')
p.runs[0].bold = True

add_code(doc,
'''results = collection.query(
    query_embeddings = [query_vector],
    n_results = 5,
    where = {
        "$and": [
            {"tenant_id": {"$eq": "company_abc"}},
            {"$or": [
                {"org_id": {"$eq": "0102"}},
                {"org_id": {"$eq": ""}},   # 전사 공유 포함
            ]}
        ]
    }
)
# 반환 결과에 metadatas 자동 포함 → source_url, page_no 등 출처 정보 제공''')

add_quote(doc, '속성 구조 상세 → §1.1 이중 축 원칙,  전체 속성 목록 → §3.2 업로드 요청 필드')

out = r'e:\ontology_edu\X_rag_std\메타데이터_관리_섹션_v4.docx'
doc.save(out)
print(f'저장: {out}')
