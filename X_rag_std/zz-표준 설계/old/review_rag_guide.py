#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Review RAG 개발 가이드_v1.1.docx"""

from docx import Document
from docx.oxml.ns import qn
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'

try:
    doc = Document(guide_file)
except Exception as e:
    print(f"파일 읽기 오류: {e}")
    sys.exit(1)

print("=" * 80)
print("RAG 개발 가이드 v1.1 검토")
print("=" * 80)

print(f"\n문서 구조:")
print(f"  - 파라그래프: {len(doc.paragraphs)}")
print(f"  - 표: {len(doc.tables)}")

# 섹션 구조 분석
print(f"\n섹션 구조 (Heading 기반):")
for i, para in enumerate(doc.paragraphs):
    if para.style.name in ['Heading 1', 'Heading 2', 'Heading 3']:
        level = para.style.name.replace('Heading ', '')
        indent = "  " * (int(level) - 1)
        text = para.text[:70]
        print(f"{indent}[H{level}] {text}")

# 전체 파라그래프 구조
print(f"\n전체 파라그래프 목록:")
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else "None"
    text = para.text[:70]
    if text.strip():
        print(f"  {i:2d} [{style:20s}] {text}")

# 표 목록
print(f"\n표 목록:")
for idx, table in enumerate(doc.tables):
    print(f"  표 {idx}: {len(table.rows)}행 x {len(table.columns)}열 | 스타일: {table.style.name if table.style else 'None'}")

# 스타일 분포
style_dist = {}
for para in doc.paragraphs:
    style = para.style.name if para.style else "None"
    style_dist[style] = style_dist.get(style, 0) + 1

print(f"\n스타일 분포:")
for style, count in sorted(style_dist.items(), key=lambda x: -x[1]):
    print(f"  {style:25s}: {count:2d}개")

# 표 포맷 검증
print(f"\n표 포맷 검증 (처음 3개):")
for t_idx in range(min(3, len(doc.tables))):
    table = doc.tables[t_idx]
    if len(table.rows) > 0:
        header_cell = table.rows[0].cells[0]
        tcPr = header_cell._element.tcPr
        has_color = False
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                has_color = fill

        border_color = "?"
        try:
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is not None:
                tblBorders = tblPr.find(qn('w:tblBorders'))
                if tblBorders is not None:
                    top_border = tblBorders.find(qn('w:top'))
                    if top_border is not None:
                        border_color = top_border.get(qn('w:color'))
        except:
            pass

        print(f"  표 {t_idx}: 헤더색={has_color if has_color else 'X'} | 테두리={border_color}")

print("\n=" * 80)
print("검토 완료!")
print("=" * 80)

