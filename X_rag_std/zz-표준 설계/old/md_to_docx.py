import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

MD_FILE = r"e:\ontology_edu\X_rag_std\RAG_표준_설계_v1.5_보고용.md"
DOCX_FILE = r"e:\ontology_edu\X_rag_std\RAG_표준_설계_v1.5_보고용.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_inline_formatting(para, text):
    """Handle bold, italic, inline code, links in a paragraph."""
    # Pattern: **bold**, `code`, [text](url), plain
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4F)
        elif re.match(r'\[([^\]]+)\]\(([^)]+)\)', part):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            run = para.add_run(m.group(1))
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
        else:
            if part:
                para.add_run(part)

def add_table_from_lines(doc, table_lines):
    # Filter out separator lines (|---|---|)
    rows = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|', l)]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        parsed.append(cells)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = 'Table Grid'
    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row):
            if ci >= cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            if ri == 0:
                run = para.add_run(clean)
                run.bold = True
                set_cell_bg(cell, 'D9E1F2')
            else:
                para.add_run(clean)
    doc.add_paragraph()

def process_md(md_file, docx_file):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Base font
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    lines = [l.rstrip('\n') for l in lines]
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if line.strip() == '':
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', m.group(2))
            text = re.sub(r'`([^`]+)`', r'\1', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F5F5F5')
            pPr.append(shd)
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            continue

        # Inline code block (indented)
        if line.startswith('  ```') or line.startswith('   ```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            continue

        # Table
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_lines(doc, table_lines)
            continue

        # Blockquote
        if line.startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            # Collect multi-line blockquote
            bq_lines = [text]
            while i + 1 < len(lines) and lines[i+1].startswith('>'):
                i += 1
                bq_lines.append(re.sub(r'^>\s*', '', lines[i]))
            full_text = ' '.join(bq_lines)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', full_text)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            run = para.add_run(clean)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.italic = True
            i += 1
            continue

        # Bullet list (*, -)
        m = re.match(r'^(\s*)[*\-]\s+(.*)', line)
        if m:
            indent = len(m.group(1))
            text = m.group(2)
            level = 'List Bullet' if indent == 0 else 'List Bullet 2'
            para = doc.add_paragraph(style=level)
            add_inline_formatting(para, text)
            i += 1
            continue

        # Numbered list (1. 2. ...)
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            para = doc.add_paragraph(style='List Number')
            add_inline_formatting(para, m.group(1))
            i += 1
            continue

        # Indented table (inside list)
        if line.startswith('  |'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_lines(doc, table_lines)
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        add_inline_formatting(para, line)
        i += 1

    doc.save(docx_file)
    print(f"저장 완료: {docx_file}")

process_md(MD_FILE, DOCX_FILE)
