#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Create final DOCX from template with all API specs"""

from docx import Document
from docx.shared import Pt, Inches
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

template_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
output_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'

# 템플릿 로드
print(f"템플릿 로드: {template_file}")
doc = Document(template_file)

# 스타일 레퍼런스 수집
heading3_style = None
normal_style = None
for para in doc.paragraphs:
    if para.style.name == 'Heading 3':
        heading3_style = para.style
    if para.style.name == 'Normal':
        normal_style = para.style

print(f"템플릿 로드 완료: {len(doc.paragraphs)} 파라그래프, {len(doc.tables)} 표")

# 3.2.1 이후 콘텐츠 제거 (재구성하기 위해)
# 원본 템플릿에서 "3.2.1 문서업로드"부터의 콘텐츠 찾기
remove_from_idx = None
for i, para in enumerate(doc.paragraphs):
    if '3.2.1' in para.text:
        remove_from_idx = i
        break

print(f"\n기존 API 섹션 제거 대기... (인덱스 {remove_from_idx}부터)")

# 마크다운에서 섹션 3.2 전체 추출
md_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md'
with open(md_file, 'r', encoding='utf-8') as f:
    md_text = f.read()

# 섹션 3.2 추출 (3.2.1부터 시작)
start = md_text.find('### 3.2.1')
end = md_text.find('## 4. 청킹 표준')
section_3_2 = md_text[start:end] if start != -1 and end != -1 else ""

if not section_3_2:
    print("ERROR: 마크다운에서 섹션 3.2를 찾을 수 없습니다.")
    sys.exit(1)

print(f"마크다운 섹션 3.2 추출: {len(section_3_2)} 문자")

# 간단한 방법: convert_to_word.py처럼 마크다운을 파싱하여 추가
lines = section_3_2.split('\n')

print(f"\n섹션 3.2 재구성 중 ({len(lines)} 라인)...")

# 기존 섹션 3.2.1 이후의 모든 내용 제거
if remove_from_idx is not None:
    # 파라그래프 역순 제거 (인덱스 변화 방지)
    for i in range(len(doc.paragraphs) - 1, remove_from_idx - 1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    print(f"파라그래프 {remove_from_idx}부터 제거 완료")

# 테이블도 제거 (3.2.1 이후의 테이블들)
# 먼저 템플릿의 테이블 개수를 세기
initial_table_count = len(doc.tables)
print(f"초기 표 개수: {initial_table_count}")

# 역순으로 제거 (3.2.1 이후 테이블)
# 표 0: API 목록
# 표 1-4: 3.2.1과 3.2.2의 내용
# 표 5+: 3.3 청킹 표준
if initial_table_count > 5:
    for i in range(initial_table_count - 1, 4, -1):
        tbl = doc.tables[i]._element
        tbl.getparent().remove(tbl)
    print(f"표 5부터 {initial_table_count-1}까지 제거 완료")

# 이제 마크다운의 섹션 3.2를 추가
# convert_to_word.py의 로직을 사용
i = 0
while i < len(lines):
    line = lines[i]

    # 제목
    if line.startswith('### '):
        p = doc.add_paragraph(line[4:], style='Heading 3')
    # 표
    elif line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|'):
        headers = [h.strip() for h in line.split('|')[1:-1]]
        i += 2  # 분리선 건너뛰기

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            if j < len(hdr_cells):
                hdr_cells[j].text = header

        while i < len(lines) and lines[i].startswith('|'):
            cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            # 헤더 개수에 맞추기
            while len(cells) < len(headers):
                cells.append('')
            row_cells = table.add_row().cells
            for j, cell in enumerate(cells[:len(headers)]):
                if j < len(row_cells):
                    row_cells[j].text = cell
            i += 1
        i -= 1
    # 코드 블록
    elif '```' in line:
        i += 1
        code_lines = []
        while i < len(lines) and '```' not in lines[i]:
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.style = 'Normal'
            p.add_run('\n'.join(code_lines)).font.name = 'Courier New'
    # 일반 텍스트
    elif line.strip() and not line.startswith('---'):
        doc.add_paragraph(line)

    i += 1

print(f"\n문서 저장 중: {output_file}")
doc.save(output_file)

print(f"✓ 완료!")
print(f"  - 파라그래프: {len(doc.paragraphs)}")
print(f"  - 표: {len(doc.tables)}")

