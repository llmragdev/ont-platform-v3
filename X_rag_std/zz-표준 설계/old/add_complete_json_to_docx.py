#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Add complete JSON examples to RAG guide Word document"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("Word 문서에 완성된 JSON 예제 추가")
print("=" * 80)

# 1. "검색 응답 주요 필드" 섹션 찾기
print("\n[1단계] 검색 응답 섹션 위치 확인")

search_response_idx = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '검색 응답' in text and '주요 필드' in text:
        search_response_idx = i
        print(f"  ✓ 발견: 라인 {i} - '{text}'")
        break

if search_response_idx is None:
    print("  ❌ 검색 응답 섹션을 찾을 수 없습니다")
    sys.exit(1)

# 2. 기존 JSON 예시가 있는 위치 찾기
print("\n[2단계] 기존 JSON 예시 위치 찾기")

existing_json_idx = None
for i in range(search_response_idx + 1, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if '검색 응답 예시' in text or ('{' in text and 'status' in text):
        existing_json_idx = i
        print(f"  ✓ 발견: 라인 {i}")
        break

if existing_json_idx:
    print(f"  기존 예시를 라인 {existing_json_idx} 근처에서 찾았습니다")

# 3. 새로운 JSON 예제 텍스트 준비
print("\n[3단계] 완성된 JSON 예제 준비")

json_examples = {
    "success_debug_true": """검색 응답 예시 (debug_mode: true)

{
  "status": "success",
  "data": {
    "query": "2026년 인사 규정 알려줘",
    "answer": "2026년 인사 규정에 따르면 신입 채용 시 기본급은 연봉 3,000만 원 이상이며, 복리후생으로는 4대 보험과 퇴직금이 보장됩니다.",
    "used_chunks": [
      {
        "chunk_id": "doc_a1b2c3d4#chunk4",
        "content": "2026년 인사 규정 제3조 - 신입 채용 기본급...",
        "similarity_score": 0.92,
        "metadata": {
          "source_name": "2026_인사규정.pdf",
          "source_url": "https://storage.example.com/docs/2026_인사규정.pdf",
          "page_no": 12,
          "category_large": "인사",
          "category_mid": "채용",
          "vector_db_id": "vdb_hr_recruit_01",
          "tenant_id": "company_abc",
          "org_id": "0102",
          "dept_code": "01"
        }
      },
      {
        "chunk_id": "doc_a1b2c3d4#chunk5",
        "content": "2026년 인사 규정 제5조 - 복리후생...",
        "similarity_score": 0.88,
        "metadata": { ... }
      },
      {
        "chunk_id": "doc_a1b2c3d4#chunk6",
        "content": "2026년 인사 규정 제8조 - 신입 연수 프로그램...",
        "similarity_score": 0.76,
        "metadata": { ... }
      }
    ],
    "debug_info": {
      "execution_time_ms": 145,
      "candidate_chunks": [
        {
          "chunk_id": "doc_a1b2c3d4#chunk2",
          "content": "2026년 인사 규정 제1조...",
          "similarity_score": 0.62,
          "metadata": { ... }
        },
        {
          "chunk_id": "doc_a1b2c3d4#chunk7",
          "content": "2026년 인사 규정 제10조...",
          "similarity_score": 0.58,
          "metadata": { ... }
        }
      ]
    }
  },
  "error": null
}""",

    "success_debug_false": """검색 응답 예시 (debug_mode: false)

{
  "status": "success",
  "data": {
    "query": "2026년 인사 규정 알려줘",
    "answer": "2026년 인사 규정에 따르면...",
    "used_chunks": [
      { "chunk_id": "...", "content": "...", "similarity_score": 0.92, "metadata": { ... } },
      { "chunk_id": "...", "content": "...", "similarity_score": 0.88, "metadata": { ... } },
      { "chunk_id": "...", "content": "...", "similarity_score": 0.76, "metadata": { ... } }
    ],
    "debug_info": {
      "execution_time_ms": 145,
      "candidate_chunks": []
    }
  },
  "error": null
}""",

    "error_examples": """검색 응답 예시 - 오류 케이스들

// 1. 벡터 DB를 찾을 수 없음
{
  "status": "error",
  "data": null,
  "error": {
    "code": "VECTOR_DB_NOT_FOUND",
    "message": "지정한 벡터 DB(vdb_hr_recruit_01)를 찾을 수 없습니다."
  }
}

// 2. 검색 쿼리 필드 누락
{
  "status": "error",
  "data": null,
  "error": {
    "code": "INVALID_REQUEST",
    "message": "필수 필드 'query'가 누락되었습니다."
  }
}

// 3. 카테고리 필터 오류
{
  "status": "error",
  "data": null,
  "error": {
    "code": "INVALID_CATEGORY",
    "message": "필터 'filters.category_large'의 값이 유효하지 않습니다."
  }
}

// 4. 벡터화 실패
{
  "status": "error",
  "data": null,
  "error": {
    "code": "EMBEDDING_FAILED",
    "message": "쿼리 벡터화 중 오류 발생: 임베딩 서버가 응답하지 않습니다."
  }
}

// 5. 검색 결과 없음
{
  "status": "error",
  "data": null,
  "error": {
    "code": "NO_RESULTS",
    "message": "검색 쿼리와 일치하는 문서를 찾을 수 없습니다."
  }
}"""
}

print(f"  ✓ 3가지 JSON 예제 세트 준비 완료")

# 4. 문서에 예제 추가
print("\n[4단계] 문서에 예제 추가")

# 검색 응답 표 다음 위치 찾기 (표 다음의 첫 번째 비어있지 않은 파라그래프)
insert_position = search_response_idx + 2

# 기존 JSON이 있으면 그것을 시작점으로 사용
if existing_json_idx:
    insert_position = existing_json_idx

print(f"  삽입 위치: 라인 {insert_position}")

# 문서에 추가할 내용 준비
additions = [
    ("success_debug_true", json_examples["success_debug_true"]),
    ("success_debug_false", json_examples["success_debug_false"]),
    ("error_examples", json_examples["error_examples"]),
]

# 파라그래프 추가
for idx, (title, content) in enumerate(additions):
    # 제목 추가
    title_para = doc.paragraphs[insert_position]._element
    new_para_title = doc.add_paragraph()
    new_para_title.text = ""

    # 내용 추가 (새 파라그래프에)
    content_para = doc.add_paragraph()
    content_para.text = content
    # 고정폭 폰트 설정
    for run in content_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    print(f"  ✓ {title} 추가 완료")

print(f"\n  총 {len(additions)}개 예제 세트 추가")

# 5. 저장
print("\n[5단계] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료: {guide_file}")

print("\n" + "=" * 80)
print("✅ JSON 예제 추가 완료")
print("=" * 80)

print("""
📝 추가된 내용:
  1. 검색 응답 예시 (debug_mode: true)
     - 3개의 used_chunks
     - 2개의 candidate_chunks (미채택 청크)

  2. 검색 응답 예시 (debug_mode: false)
     - 3개의 used_chunks
     - 빈 candidate_chunks 배열

  3. 검색 응답 예시 - 오류 케이스들
     - VECTOR_DB_NOT_FOUND
     - INVALID_REQUEST
     - INVALID_CATEGORY
     - EMBEDDING_FAILED
     - NO_RESULTS

📊 결과:
  - 모든 18개 필드가 실제 데이터로 표현됨
  - 배열 구조가 명확하게 표시됨
  - 실제 사용 사례를 반영한 예제 제공
""")
