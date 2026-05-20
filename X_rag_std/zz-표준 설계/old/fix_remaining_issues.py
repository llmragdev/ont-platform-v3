#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix remaining issues - typos and duplicate section"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

fixes = []

print("=" * 80)
print("남은 문제 수정")
print("=" * 80)

# 1. 백터DB 오타 수정 (모든 occurrence)
print("\n[1단계] 백터DB → 벡터DB 오타 수정")

typo_count = 0
for i, para in enumerate(doc.paragraphs):
    if '백터DB' in para.text:
        for run in para.runs:
            if '백터DB' in run.text:
                old_text = run.text
                run.text = run.text.replace('백터DB', '벡터DB')
                typo_count += 1
                fixes.append({
                    'type': '오타 수정',
                    'line': i,
                    'from': '백터DB',
                    'to': '벡터DB'
                })
                print(f"  ✓ 라인 {i}: '백터DB' → '벡터DB'")

# 2. 라인 147의 중복 섹션을 3.2.5 문서삭제로 변경
print("\n[2단계] 중복 섹션 제목 수정 (3.2.5 문서삭제)")

# 라인 147 근처에서 DELETE가 포함된 "3.2.4 문서재업로드"를 찾아서 "3.2.5 문서삭제"로 변경
found_delete = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # 3.2.4 또는 문서재업로드이면서 DELETE 포함
    if 'DELETE' in text and ('3.2.4' in text or '문서재업로드' in text):
        if i > 140:  # 라인 147 근처
            old_text = para.text.strip()
            # 기존 텍스트를 새 텍스트로 교체
            para.clear()
            if '3.2.4' in old_text:
                new_text = old_text.replace('3.2.4 문서재업로드', '3.2.5 문서삭제')
            else:
                new_text = '3.2.5 문서삭제 — `DELETE /api/v1/documents/{doc_id}`'
            para.add_run(new_text)
            fixes.append({
                'type': '중복 섹션 수정',
                'line': i,
                'from': old_text[:70],
                'to': new_text[:70]
            })
            print(f"  ✓ 라인 {i}: '3.2.4 문서재업로드' → '3.2.5 문서삭제'")
            found_delete = True
            break

if not found_delete:
    # 직접 찾기
    for i in range(140, min(160, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if 'DELETE' in text or 'deleted_chunks' in text:
            if '3.2.4' in text or '문서재업로드' in text:
                old_text = doc.paragraphs[i].text
                doc.paragraphs[i].clear()
                new_text = '3.2.5 문서삭제 — `DELETE /api/v1/documents/{doc_id}`'
                doc.paragraphs[i].add_run(new_text)
                fixes.append({
                    'type': '중복 섹션 수정',
                    'line': i,
                    'from': old_text.strip()[:70],
                    'to': new_text[:70]
                })
                print(f"  ✓ 라인 {i}: → '3.2.5 문서삭제'")
                break

# 저장
print(f"\n[최종] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료")

# 요약
print("\n" + "=" * 80)
print("남은 문제 수정 완료")
print("=" * 80)

print(f"\n총 {len(fixes)}개 항목 추가 수정:")
for idx, fix in enumerate(fixes, 1):
    print(f"\n[{idx}] {fix['type']}")
    print(f"    라인: {fix['line']}")
    print(f"    변경: {fix['from']} → {fix['to']}")

print("\n" + "=" * 80)

