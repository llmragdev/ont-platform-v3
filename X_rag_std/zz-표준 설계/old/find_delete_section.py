#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Find DELETE section and verify"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("DELETE 섹션 찾기")
print("=" * 80)

# 라인 145-160 확인
print("\n라인 145-160 내용:")
for i in range(145, min(160, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    if text:
        print(f"  라인 {i}: {text[:80]}")

# DELETE 또는 deleted_chunks 포함 라인 찾기
print("\n\nDELETE/삭제 관련 라인:")
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if 'DELETE' in text or 'deleted_chunks' in text or '문서삭제' in text:
        print(f"  라인 {i}: {text.strip()[:80]}")

# 3.2.4가 2번 나타나는지 확인
print("\n\n3.2.4 섹션 확인:")
api_4_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.2.4' in text:
        api_4_count += 1
        print(f"  라인 {i}: {text[:80]}")

print(f"\n3.2.4가 {api_4_count}번 나타남")

if api_4_count == 2:
    print("⚠️ 3.2.4가 2번 나타남 - 하나를 3.2.5로 변경해야 함")

# 3.2.5 확인
print("\n3.2.5 섹션 확인:")
api_5_found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.2.5' in text:
        api_5_found = True
        print(f"  라인 {i}: {text[:80]}")

if not api_5_found:
    print("  ❌ 3.2.5가 없음 - 추가해야 함")

