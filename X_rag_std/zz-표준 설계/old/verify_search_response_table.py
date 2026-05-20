#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Verify search response table and apply final formatting"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("검색 응답 표 검증 및 최종 포맷팅")
print("=" * 80)

# 첫 번째 표 (검색 응답 표)
if len(doc.tables) > 0:
    table = doc.tables[0]
    print(f"\n표 정보:")
    print(f"  행: {len(table.rows)}")
    print(f"  열: {len(table.rows[0].cells) if len(table.rows) > 0 else 0}")

    # 표의 테두리를 검정색으로 설정
    print(f"\n[포맷팅] 표 테두리를 검정색(000000)으로 설정")

    def set_cell_border(cell, **kwargs):
        """
        테이블 셀의 테두리 설정
        """
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '12')  # 1.5pt
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')  # 검정색
            tcBorders.append(edge_element)

        tcPr.append(tcBorders)

    # 모든 셀에 검정 테두리 적용
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)

    print(f"  ✓ 모든 셀 {len(table.rows) * len(table.rows[0].cells)}개에 적용")

    # 표 콘텐츠 확인
    print(f"\n표 내용 확인:")
    for i, row in enumerate(table.rows):
        cells_text = " | ".join([cell.text[:30] for cell in row.cells])
        if i == 0:
            print(f"  [헤더] {cells_text}")
        elif i <= 5:
            print(f"  [{i:2d}] {cells_text}")
    if len(table.rows) > 6:
        print(f"  ... (총 {len(table.rows)}행)")

# 저장
print(f"\n[최종] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료: {guide_file}")

print("\n" + "=" * 80)
print("✓ 검색 응답 표 최종 포맷팅 완료!")
print("=" * 80)
print("\n📊 업데이트 요약:")
print(f"  - 필드 수: 8개 → 18개 (10개 필드 추가)")
print(f"  - 구조: 기본 필드 → 완전한 JSON 명세")
print(f"  - 테두리: 검정색(000000)")
print(f"  - 헤더: 파란색(D9E1F2) 배경")
