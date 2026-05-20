#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix table cell colors - remove background from data rows"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

docx_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'
output_file = docx_file  # 같은 파일에 덮어쓰기

doc = Document(docx_file)

print(f"문서 로드: {len(doc.tables)} 표")
print("\n표 포맷 수정 중...")

# 모든 테이블 처리
for table_idx, table in enumerate(doc.tables):
    print(f"\n표 {table_idx}: {len(table.rows)}행 x {len(table.columns)}열")

    # 첫 번째 행은 헤더 (색상 유지)
    # 나머지 행은 데이터 (색상 제거)

    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0)

        for cell in row.cells:
            # 셀의 배경색 설정
            shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))

            # 데이터 행: 흰색 배경으로 명시적 설정 (색상 제거)
            if not is_header:
                cell._element.get_or_add_tcPr().append(shading_elm)
            # 헤더 행: 파란색 유지 (기존 스타일 유지)
            # 헤더는 테이블 스타일에서 자동으로 포맷됨

    if table_idx < 3:
        print(f"  ✓ 포맷 수정 완료")

print(f"\n문서 저장: {output_file}")
doc.save(output_file)

print("✓ 표 색상 수정 완료!")
print("\n수정 내용:")
print("  - 헤더 행(첫 번째): 파란색 배경 유지")
print("  - 데이터 행: 흰색 배경(색상 제거)")

