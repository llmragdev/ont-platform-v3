#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

template = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
doc = Document(template)

print("=" * 80)
print("사용자 편집본 (_내가편집.docx) 전체 내용")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else "None"
    text = para.text[:70]
    print(f"{i:2d} [{style:20s}] {text}")

print("\n" + "=" * 80)
print("표 내용")
print("=" * 80)

for t_idx, table in enumerate(doc.tables):
    print(f"\n표 {t_idx}: {len(table.rows)}행 x {len(table.columns)}열")
    for r_idx, row in enumerate(table.rows[:3]):
        cells = [cell.text[:15] for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")

print("\n완료!")
