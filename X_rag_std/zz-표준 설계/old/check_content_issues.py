#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Check RAG guide for content issues"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("RAG 개발 가이드 내용 검토")
print("=" * 80)

issues = []

# 1. 중복/오류 섹션 찾기
print("\n[1] 중복/유사 섹션 확인")
api_sections = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '문서' in text and ('POST' in text or 'GET' in text or 'PUT' in text or 'DELETE' in text):
        api_sections.append((i, text))

print(f"\nAPI 섹션 목록:")
for line, text in api_sections:
    print(f"  라인 {line:3d}: {text[:80]}")

# 중복 확인
for i in range(len(api_sections)-1):
    for j in range(i+1, len(api_sections)):
        if '문서재업로드' in api_sections[i][1] and '문서재업로드' in api_sections[j][1]:
            issues.append({
                'type': '❌ 중복 섹션',
                'line1': api_sections[i][0],
                'line2': api_sections[j][0],
                'text': '문서재업로드 (2번 반복됨)',
                'note': '두 번째는 "문서삭제 — DELETE"로 변경되어야 함'
            })

# 2. 문서 버전 확인
print("\n[2] 문서 버전/날짜 확인")
for i, para in enumerate(doc.paragraphs[:10]):
    text = para.text.strip()
    if 'v1.0' in text or 'v1.1' in text or '2026' in text:
        print(f"  라인 {i}: {text}")
        if 'v1.0' in text and 'v1.1' not in text:
            issues.append({
                'type': '⚠️ 버전 불일치',
                'line': i,
                'text': text,
                'note': '파일명은 v1.1인데 내용은 v1.0'
            })

# 3. 상태 코드 확인 (pending → 01 등으로 변경되었는지)
print("\n[3] 상태 코드 확인")
status_code_issues = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'pending' in text or 'processing' in text or 'completed' in text or 'error' in text.lower():
        if not any(code in text for code in ['01', '02', '03', '04']):
            if 'error_code' not in text.lower() and 'error' in text.lower():
                continue
            status_code_issues.append((i, text))

if status_code_issues:
    print(f"\n  구형 상태 코드 발견:")
    for line, text in status_code_issues[:5]:
        print(f"    라인 {line}: {text[:70]}")
    if len(status_code_issues) > 5:
        print(f"    ... 외 {len(status_code_issues)-5}개")

# 4. 오타/맞춤법 확인
print("\n[4] 주요 오타/맞춤법")
typo_patterns = [
    ('백터DB', '벡터DB', 'DB 철자'),
    ('케테고리', '카테고리', '카테고리 철자'),
]

typos_found = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    for wrong, correct, desc in typo_patterns:
        if wrong in text:
            typos_found.append({
                'line': i,
                'text': text[:70],
                'wrong': wrong,
                'correct': correct,
                'desc': desc
            })

if typos_found:
    print(f"\n  오타 발견:")
    for typo in typos_found[:5]:
        print(f"    라인 {typo['line']}: '{typo['wrong']}' → '{typo['correct']}' ({typo['desc']})")
        issues.append({
            'type': '🔤 오타',
            'line': typo['line'],
            'from': typo['wrong'],
            'to': typo['correct'],
            'desc': typo['desc']
        })

# 5. 내용 일관성 확인
print("\n[5] 내용 일관성 확인")

# 섹션 3.2의 API 개수 확인
section3_2_apis = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('3.2.'):
        section3_2_apis.append((i, text))

print(f"\n  섹션 3.2 API 스펙:")
for line, text in section3_2_apis:
    print(f"    라인 {line}: {text[:60]}")

# 3.2.1 ~ 3.2.5가 모두 있는지 확인
expected_apis = ['3.2.1', '3.2.2', '3.2.3', '3.2.4', '3.2.5']
found_apis = [text for _, text in section3_2_apis]
missing = [api for api in expected_apis if not any(api in text for text in found_apis)]

if missing:
    issues.append({
        'type': '❌ 누락된 API',
        'missing': missing,
        'note': f'3.2 섹션에서 {missing}이(가) 누락됨'
    })
else:
    print(f"  ✓ 모든 API (3.2.1~3.2.5) 포함됨")

# 6. 정보 최신성 확인
print("\n[6] 정보 최신성")
outdated_terms = [
    '예정',
    '임시',
    '보류',
    'TODO',
]

for i, para in enumerate(doc.paragraphs):
    text = para.text
    for term in outdated_terms:
        if term in text and ('보류' not in text or '접근 권한 관리' not in text):
            if i < 240:  # 섹션 7 전까지만
                print(f"  라인 {i}: '{term}' 포함 - {text[:60]}")
                break

# ========== 결과 출력 ==========
print("\n" + "=" * 80)
print("내용 검토 결과")
print("=" * 80)

if issues:
    print(f"\n수정 필요 사항 ({len(issues)}개):\n")
    for idx, issue in enumerate(issues, 1):
        print(f"[{idx}] {issue.get('type', '기타')}")
        if 'line' in issue:
            print(f"    위치: 라인 {issue['line']}")
        if 'text' in issue:
            print(f"    내용: {issue['text'][:70]}")
        if 'from' in issue and 'to' in issue:
            print(f"    변경: '{issue['from']}' → '{issue['to']}'")
        if 'note' in issue:
            print(f"    참고: {issue['note']}")
        if 'missing' in issue:
            print(f"    누락: {issue['missing']}")
        print()
else:
    print("\n✓ 내용상 주요 오류 없음")

print("=" * 80)

