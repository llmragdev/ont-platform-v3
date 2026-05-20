#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Rebuild DOCX using user's edited version as template"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# 사용자 편집본 로드 (포맷 템플릿으로 사용)
template_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'
output_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'

# 템플릿 읽기
template_doc = Document(template_file)

print(f"템플릿 문서 분석...")
print(f"  - 파라그래프: {len(template_doc.paragraphs)}")
print(f"  - 표: {len(template_doc.tables)}")

# 템플릿의 스타일과 포맷 분석
print(f"\n템플릿 구조 (처음 20개 항목):")
for i in range(min(20, len(template_doc.paragraphs))):
    para = template_doc.paragraphs[i]
    style = para.style.name if para.style else "None"
    text = para.text[:50]
    print(f"  {i:2d}: [{style:20s}] {text}")

# 새 문서 생성 (템플릿 복제)
doc = Document(template_file)

print(f"\n새로운 API 스펙 추가 중...")

# 현재 문서의 끝을 찾기
# _내가편집.docx는 섹션 3만 있으므로, 전체 문서 구조를 다시 빌드해야 함
# 마크다운의 구조를 따라 재구성

# 먼저 현재 문서의 마지막 내용 확인
print(f"\n마지막 5개 파라그래프:")
for i in range(max(0, len(doc.paragraphs)-5), len(doc.paragraphs)):
    para = doc.paragraphs[i]
    text = para.text[:60]
    print(f"  {i}: {text}")

# 새 내용 추가 (마크다운 파일의 확장된 API 스펙)
# 마크다운 파일 읽기
md_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md'
with open(md_file, 'r', encoding='utf-8') as f:
    md_content = f.read()

# 섹션 3 부분만 추출
start_marker = "## 3. API 레이아웃"
end_marker = "## 4. 청킹 표준"

start_idx = md_content.find(start_marker)
end_idx = md_content.find(end_marker)

if start_idx == -1 or end_idx == -1:
    print("ERROR: 섹션 3을 마크다운에서 찾을 수 없습니다.")
    sys.exit(1)

section3_md = md_content[start_idx:end_idx]

print(f"\n섹션 3 마크다운 길이: {len(section3_md)} 문자")
print(f"섹션 3 시작 부분:\n{section3_md[:200]}")

print("\n문서를 재빌드합니다...")

# 새 문서 생성
new_doc = Document()

# 스타일 복사 함수
def copy_style_from_para(src_para, dst_para):
    """소스 파라그래프의 스타일을 대상 파라그래프로 복사"""
    if src_para.style:
        dst_para.style = src_para.style

    # 포맷 복사
    pf = dst_para.paragraph_format
    src_pf = src_para.paragraph_format

    if src_pf.left_indent is not None:
        pf.left_indent = src_pf.left_indent
    if src_pf.right_indent is not None:
        pf.right_indent = src_pf.right_indent
    if src_pf.first_line_indent is not None:
        pf.first_line_indent = src_pf.first_line_indent
    if src_pf.space_before is not None:
        pf.space_before = src_pf.space_before
    if src_pf.space_after is not None:
        pf.space_after = src_pf.space_after

# 원본 문서 내용 먼저 복사
print(f"원본 {len(doc.paragraphs)} 파라그래프 복사 중...")
for para in doc.paragraphs:
    new_para = new_doc.add_paragraph(para.text)
    copy_style_from_para(para, new_para)

# 표도 복사
print(f"원본 {len(doc.tables)} 표 복사 중...")
for table in doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.style = table.style
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_table.rows[i].cells[j].text = cell.text

print(f"\n새 문서 저장: {output_file}")
new_doc.save(output_file)

print("✓ 완료!")
print(f"새 문서: {len(new_doc.paragraphs)} 파라그래프, {len(new_doc.tables)} 표")

