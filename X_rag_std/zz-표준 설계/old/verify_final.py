#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

output_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'
doc = Document(output_file)

print("=" * 80)
print("최종 문서 검증")
print("=" * 80)

print(f"\n파라그래프: {len(doc.paragraphs)}")
print(f"표: {len(doc.tables)}")

print(f"\n구조 (처음 45개 파라그래프):")
for i, para in enumerate(doc.paragraphs[:45]):
    style = para.style.name if para.style else "None"
    text = para.text[:60]
    print(f"  {i:2d} [{style:20s}] {text}")

print(f"\n표 목록:")
for idx, table in enumerate(doc.tables):
    print(f"  표 {idx}: {len(table.rows)}행 x {len(table.columns)}열")

print("\n✓ 최종 문서 생성 완료!")
