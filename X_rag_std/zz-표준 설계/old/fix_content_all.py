#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix all content issues in RAG 개발 가이드"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

fixes = []

print("=" * 80)
print("RAG 개발 가이드 내용 수정 시작")
print("=" * 80)

# 1. 버전 수정 (v1.0 → v1.1)
print("\n[1단계] 버전 불일치 수정")
if doc.paragraphs[0].text.strip() == "RAG 개발 설계 가이드 v1.0":
    old_text = doc.paragraphs[0].text
    doc.paragraphs[0].clear()
    doc.paragraphs[0].add_run("RAG 개발 설계 가이드 v1.1")
    fixes.append({
        'type': '버전 수정',
        'line': 0,
        'from': old_text.strip(),
        'to': 'RAG 개발 설계 가이드 v1.1'
    })
    print(f"  ✓ 라인 0: v1.0 → v1.1")

# 2. 오타 수정
print("\n[2단계] 오타 수정")

typos = [
    ('케테고리', '카테고리'),
    ('백터DB', '벡터DB'),
]

for wrong, correct in typos:
    for i, para in enumerate(doc.paragraphs):
        if wrong in para.text:
            # 정확히 텍스트 교체
            for run in para.runs:
                if wrong in run.text:
                    old_run_text = run.text
                    run.text = run.text.replace(wrong, correct)
                    fixes.append({
                        'type': '오타 수정',
                        'line': i,
                        'from': wrong,
                        'to': correct,
                        'text_context': para.text[:70]
                    })
                    print(f"  ✓ 라인 {i}: '{wrong}' → '{correct}'")

# 3. 중복 섹션 수정 (문서재업로드 → 문서삭제)
print("\n[3단계] 중복 섹션 수정")

api_sections_found = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '문서재업로드' in text and 'PUT' in text:
        api_sections_found.append((i, text))

if len(api_sections_found) >= 2:
    # 두 번째 "문서재업로드"를 "문서삭제"로 변경
    second_idx = api_sections_found[1][0]
    old_text = doc.paragraphs[second_idx].text.strip()

    # DELETE 엔드포인트인지 확인
    is_delete = 'DELETE' in old_text

    if is_delete:
        new_text = old_text.replace('문서재업로드', '문서삭제')
        doc.paragraphs[second_idx].clear()
        doc.paragraphs[second_idx].add_run(new_text)
        fixes.append({
            'type': '중복 섹션 수정',
            'line': second_idx,
            'from': old_text,
            'to': new_text
        })
        print(f"  ✓ 라인 {second_idx}: '문서재업로드' → '문서삭제'")

# 4. API 섹션 헤더 명시화 (3.2.1, 3.2.2, 3.2.3, 3.2.4, 3.2.5)
print("\n[4단계] API 섹션 헤더 명시화")

api_mappings = [
    ('문서 업로드', '3.2.1 문서업로드'),
    ('문서목록조회', '3.2.2 문서목록조회'),
    ('문서상태조회', '3.2.3 문서상태조회'),
]

# 첫 번째 "문서재업로드"는 3.2.4
# 수정된 "문서삭제"는 3.2.5

found_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # 3.2.1
    if text.startswith('문서 업로드—') and 'POST' in text and not text.startswith('3.2.'):
        old_text = text
        new_text = text.replace('문서 업로드—', '3.2.1 문서업로드 —')
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(new_text)
        fixes.append({
            'type': 'API 헤더 명시화',
            'line': i,
            'from': old_text[:60],
            'to': new_text[:60],
            'section': '3.2.1'
        })
        print(f"  ✓ 라인 {i}: → 3.2.1 문서업로드")
        found_count += 1

    # 3.2.2
    elif text.startswith('문서목록조회 —') and 'GET' in text and not text.startswith('3.2.'):
        old_text = text
        new_text = text.replace('문서목록조회 —', '3.2.2 문서목록조회 —')
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(new_text)
        fixes.append({
            'type': 'API 헤더 명시화',
            'line': i,
            'from': old_text[:60],
            'to': new_text[:60],
            'section': '3.2.2'
        })
        print(f"  ✓ 라인 {i}: → 3.2.2 문서목록조회")
        found_count += 1

    # 3.2.3
    elif text.startswith('문서상태조회 —') and 'GET' in text and not text.startswith('3.2.'):
        old_text = text
        new_text = text.replace('문서상태조회 —', '3.2.3 문서상태조회 —')
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(new_text)
        fixes.append({
            'type': 'API 헤더 명시화',
            'line': i,
            'from': old_text[:60],
            'to': new_text[:60],
            'section': '3.2.3'
        })
        print(f"  ✓ 라인 {i}: → 3.2.3 문서상태조회")
        found_count += 1

    # 3.2.4 - 첫 번째 "문서재업로드"
    elif text.startswith('문서재업로드 —') and 'PUT' in text and not text.startswith('3.2.'):
        old_text = text
        new_text = text.replace('문서재업로드 —', '3.2.4 문서재업로드 —')
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(new_text)
        fixes.append({
            'type': 'API 헤더 명시화',
            'line': i,
            'from': old_text[:60],
            'to': new_text[:60],
            'section': '3.2.4'
        })
        print(f"  ✓ 라인 {i}: → 3.2.4 문서재업로드")
        found_count += 1

    # 3.2.5 - "문서삭제" (수정된 것)
    elif text.startswith('문서삭제 —') and 'DELETE' in text and not text.startswith('3.2.'):
        old_text = text
        new_text = text.replace('문서삭제 —', '3.2.5 문서삭제 —')
        doc.paragraphs[i].clear()
        doc.paragraphs[i].add_run(new_text)
        fixes.append({
            'type': 'API 헤더 명시화',
            'line': i,
            'from': old_text[:60],
            'to': new_text[:60],
            'section': '3.2.5'
        })
        print(f"  ✓ 라인 {i}: → 3.2.5 문서삭제")
        found_count += 1

print(f"  총 {found_count}개 API 헤더 명시화 완료")

# 저장
print(f"\n[최종] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료: {guide_file}")

# 요약
print("\n" + "=" * 80)
print("수정 완료 요약")
print("=" * 80)

fix_types = {}
for fix in fixes:
    ftype = fix['type']
    fix_types[ftype] = fix_types.get(ftype, 0) + 1

print(f"\n총 {len(fixes)}개 항목 수정:\n")
for ftype, count in sorted(fix_types.items()):
    print(f"  {ftype}: {count}개")

print(f"\n상세 수정 내용:\n")
for idx, fix in enumerate(fixes, 1):
    print(f"[{idx}] {fix['type']}")
    print(f"    라인: {fix['line']}")
    if 'from' in fix and 'to' in fix:
        if len(str(fix['from'])) < 80 and len(str(fix['to'])) < 80:
            print(f"    변경: {fix['from']} → {fix['to']}")
        else:
            print(f"    변경: {str(fix['from'])[:60]}... → {str(fix['to'])[:60]}...")
    if 'section' in fix:
        print(f"    섹션: {fix['section']}")
    print()

print("=" * 80)
print("✓ 모든 수정 완료!")
print("=" * 80)

