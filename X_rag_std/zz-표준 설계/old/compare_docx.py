#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Compare two DOCX files to check formatting differences"""

from docx import Document
import sys

# stdout 인코딩을 UTF-8로 설정
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

file1 = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'
file2 = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리_내가편집.docx'

try:
    doc1 = Document(file1)
    doc2 = Document(file2)
except Exception as e:
    print(f"파일 읽기 오류: {e}")
    sys.exit(1)

print("\n" + "="*80)
print("파일 구조 비교")
print("="*80)

print(f"\n[파일1] 자동 생성 docx")
print(f"  - 파라그래프: {len(doc1.paragraphs)}")
print(f"  - 표: {len(doc1.tables)}")
print(f"  - 섹션 스타일 분석:")

para_styles = {}
for para in doc1.paragraphs:
    style = para.style.name if para.style else "None"
    para_styles[style] = para_styles.get(style, 0) + 1

for style, count in sorted(para_styles.items()):
    print(f"    {style}: {count}")

print(f"\n[파일2] 사용자 편집 docx")
print(f"  - 파라그래프: {len(doc2.paragraphs)}")
print(f"  - 표: {len(doc2.tables)}")
print(f"  - 섹션 스타일 분석:")

para_styles2 = {}
for para in doc2.paragraphs:
    style = para.style.name if para.style else "None"
    para_styles2[style] = para_styles2.get(style, 0) + 1

for style, count in sorted(para_styles2.items()):
    print(f"    {style}: {count}")

print("\n" + "="*80)
print("텍스트 내용 비교")
print("="*80)

doc1_text = '\n'.join([p.text for p in doc1.paragraphs])
doc2_text = '\n'.join([p.text for p in doc2.paragraphs])

if doc1_text == doc2_text:
    print("\n✓ 텍스트 내용은 동일합니다.")
else:
    print("\n✗ 텍스트 내용이 다릅니다.")
    print(f"  파일1 텍스트 길이: {len(doc1_text)}")
    print(f"  파일2 텍스트 길이: {len(doc2_text)}")

print("\n" + "="*80)
print("표 구조 비교")
print("="*80)

print(f"\n파일1 표 구조:")
for idx, table in enumerate(doc1.tables):
    print(f"  표 {idx}: {len(table.rows)}행 x {len(table.columns)}열")

print(f"\n파일2 표 구조:")
for idx, table in enumerate(doc2.tables):
    print(f"  표 {idx}: {len(table.rows)}행 x {len(table.columns)}열")

print("\n" + "="*80)
print("Word 포맷 스타일 비교 (첫 15개 파라그래프)")
print("="*80)

print(f"\n파일1:")
for i, para in enumerate(doc1.paragraphs[:15]):
    style = para.style.name if para.style else "None"
    indent = para.paragraph_format.left_indent
    space_before = para.paragraph_format.space_before
    space_after = para.paragraph_format.space_after
    print(f"  {i:2d}: [{style:20s}] indent={indent} space_before={space_before} space_after={space_after}")

print(f"\n파일2:")
for i, para in enumerate(doc2.paragraphs[:15]):
    style = para.style.name if para.style else "None"
    indent = para.paragraph_format.left_indent
    space_before = para.paragraph_format.space_before
    space_after = para.paragraph_format.space_after
    print(f"  {i:2d}: [{style:20s}] indent={indent} space_before={space_before} space_after={space_after}")

print("\n" + "="*80)
print("결론")
print("="*80)

# 스타일 비교
style_diff = set(para_styles.keys()) ^ set(para_styles2.keys())
if style_diff:
    print(f"\n! 사용된 스타일이 다릅니다: {style_diff}")
else:
    print("\n✓ 사용된 스타일은 같습니다.")

# 표 개수 비교
if len(doc1.tables) == len(doc2.tables):
    print("✓ 표 개수는 같습니다.")
else:
    print(f"! 표 개수가 다릅니다: {len(doc1.tables)} vs {len(doc2.tables)}")

print("\n완료!")
