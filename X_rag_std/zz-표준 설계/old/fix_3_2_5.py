#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix 3.2.5 - change duplicate 3.2.4 to 3.2.5"""

from docx import Document
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
doc = Document(guide_file)

print("=" * 80)
print("3.2.5 문서삭제 섹션 수정")
print("=" * 80)

# 라인 147: "3.2.4 문서재업로드 — `PUT ..."를 "3.2.5 문서삭제 — `DELETE ..."로 변경
target_line = 147

if target_line < len(doc.paragraphs):
    para = doc.paragraphs[target_line]
    old_text = para.text.strip()

    print(f"\n라인 {target_line}:")
    print(f"  현재: {old_text}")

    # 새 텍스트
    new_text = '3.2.5 문서삭제 — `DELETE /api/v1/documents/{doc_id}`'

    print(f"  변경: {new_text}")

    # 파라그래프 내용 변경
    style_backup = para.style
    para.clear()
    para.style = style_backup
    para.add_run(new_text)

    print(f"\n  ✓ 수정 완료")

# 저장
print(f"\n[최종] 문서 저장")
doc.save(guide_file)
print(f"  ✓ 저장 완료: {guide_file}")

print("\n" + "=" * 80)
print("✓ 모든 수정 완료!")
print("=" * 80)

