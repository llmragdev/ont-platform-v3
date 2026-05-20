#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix RAG 개발 가이드_v1.1.docx - track all changes"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import sys

if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

guide_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG 개발 가이드_v1.1.docx'
output_file = guide_file  # 같은 파일에 덮어쓰기

doc = Document(guide_file)

# 수정 사항 추적
changes = []

print("=" * 80)
print("RAG 개발 가이드 수정 시작")
print("=" * 80)

# 1. 제목 구조 수정
print("\n[1단계] 제목 구조 수정 (Heading 레벨 적용)")

# 라인 0: "RAG 개발 설계 가이드 v1.0" → Heading 1
if doc.paragraphs[0].text.strip() == "RAG 개발 설계 가이드 v1.0":
    doc.paragraphs[0].style = 'Heading 1'
    changes.append({
        'type': '제목 스타일 변경',
        'line': 0,
        'text': doc.paragraphs[0].text[:50],
        'from': 'Normal',
        'to': 'Heading 1'
    })
    print(f"  ✓ 라인 0: 'RAG 개발 설계 가이드 v1.0' → Heading 1")

# 주요 섹션 제목들을 Heading 2로 변경
heading2_texts = [
    '1. 개요 및 설계 원칙',
    '2. RAG기반 검색',
    '3. 임베딩 대상 문서 관리',
    '4. 백터DB 관리',
    '5. 백터DB의 임베딩별 문서 매칭관리',
    '6. 메타데이터 관리',
    '7. 접근 권한 관리 (보류)',
    '8. 오류/예외 결과'
]

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in heading2_texts and para.style.name != 'Heading 2':
        old_style = para.style.name
        para.style = 'Heading 2'
        changes.append({
            'type': '제목 스타일 변경',
            'line': i,
            'text': text,
            'from': old_style,
            'to': 'Heading 2'
        })
        print(f"  ✓ 라인 {i}: '{text}' → Heading 2")

# 2. 중복 섹션 수정 (3.2.4와 3.2.5)
print("\n[2단계] 중복 섹션 수정 (3.2.5 문서삭제)")

# 라인 147 근처에서 "문서재업로드"를 찾아서 "문서삭제"로 변경
for i, para in enumerate(doc.paragraphs):
    if i > 145 and i < 150:
        text = para.text.strip()
        if '문서재업로드' in text and 'DELETE' in text:
            # 이 문서는 DELETE이므로 문서삭제로 변경
            new_text = text.replace('문서재업로드', '문서삭제')
            # 파라그래프의 첫 run만 변경
            if para.runs:
                para.clear()
                para.add_run(new_text)
            changes.append({
                'type': '텍스트 수정',
                'line': i,
                'text_before': text,
                'text_after': new_text
            })
            print(f"  ✓ 라인 {i}: '문서재업로드' → '문서삭제'")
            break

# 3. List Paragraph를 Normal로 변경 (API 섹션)
print("\n[3단계] 스타일 통일 (List Paragraph → Normal)")

api_related_keywords = [
    '요청 필드',
    '응답 필드',
    '응답 예시',
    '헤더타입',
]

list_para_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # API 섹션의 List Paragraph를 Normal로 변경
    if para.style.name == 'List Paragraph' and any(kw in text for kw in api_related_keywords):
        old_style = para.style.name
        para.style = 'Normal'
        list_para_count += 1
        if list_para_count <= 10:  # 처음 10개만 상세 출력
            changes.append({
                'type': 'API 섹션 스타일 변경',
                'line': i,
                'text': text[:50],
                'from': old_style,
                'to': 'Normal'
            })

print(f"  ✓ List Paragraph → Normal로 변경: {list_para_count}개")

# 4. 테이블 테두리 색상 확인 및 설정 (검정색)
print("\n[4단계] 테이블 테두리 색상 설정")

def set_table_borders_black(table):
    """테이블의 모든 테두리를 검정색으로 설정"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr {}/>', nsdecls('w'))
        tbl.insert(0, tblPr)

    # 기존 테두리 설정 제거 후 새것 추가
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)

    # 검정색 테두리 추가
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

table_count = 0
for idx, table in enumerate(doc.tables):
    try:
        set_table_borders_black(table)
        table_count += 1
    except:
        pass

print(f"  ✓ 테이블 테두리 검정색 설정: {table_count}개/{len(doc.tables)}")
changes.append({
    'type': '테이블 테두리 색상 설정',
    'count': table_count,
    'description': '모든 테이블의 테두리를 검정색(000000)으로 설정'
})

# 5. 저장
print(f"\n[최종] 문서 저장")
doc.save(output_file)
print(f"  ✓ 저장 완료: {output_file}")

# 수정 사항 요약
print("\n" + "=" * 80)
print("수정 사항 요약")
print("=" * 80)

print(f"\n총 {len(changes)}개 항목 수정됨:\n")

change_types = {}
for change in changes:
    ctype = change.get('type', '기타')
    change_types[ctype] = change_types.get(ctype, 0) + 1

for ctype, count in sorted(change_types.items()):
    print(f"  {ctype}: {count}개")

print(f"\n" + "=" * 80)
print("상세 수정 내용")
print("=" * 80)

for idx, change in enumerate(changes, 1):
    print(f"\n[{idx}] {change.get('type', '기타')}")
    if 'line' in change:
        print(f"    위치: 라인 {change['line']}")
    if 'text' in change:
        print(f"    내용: {change.get('text', '')[:70]}")
    if 'from' in change and 'to' in change:
        print(f"    변경: {change['from']} → {change['to']}")
    if 'text_before' in change:
        print(f"    변경: '{change['text_before']}'")
        print(f"         → '{change['text_after']}'")
    if 'description' in change:
        print(f"    설명: {change['description']}")

print(f"\n" + "=" * 80)
print("✓ 수정 완료!")
print("=" * 80)

