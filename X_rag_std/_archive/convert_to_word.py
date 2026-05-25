#!/usr/bin/env python3
"""Convert RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md to DOCX"""

from docx import Document
from docx.shared import Pt, RGBColor

md_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.md'
docx_file = r'E:\ontology_edu\X_rag_std\zz-표준 설계\RAG_표준_설계_v1.5_임베딩 대상 문서 관리.docx'

# MD 파일 읽기
with open(md_file, 'r', encoding='utf-8') as f:
    content = f.read()

# Word 문서 생성
doc = Document()

lines = content.split('\n')
i = 0

while i < len(lines):
    line = lines[i]

    # 제목
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    # 표 감지
    elif line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|'):
        headers = [h.strip() for h in line.split('|')[1:-1]]
        i += 2  # 분리선 건너뛰기

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for j, header in enumerate(headers):
            hdr_cells[j].text = header

        while i < len(lines) and lines[i].startswith('|'):
            cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            # 헤더 개수에 맞춰서 셀 추가
            while len(cells) < len(headers):
                cells.append('')
            row_cells = table.add_row().cells
            for j, cell in enumerate(cells[:len(headers)]):
                row_cells[j].text = cell
            i += 1
        i -= 1
    # 코드 블록
    elif '```' in line:
        i += 1
        code_lines = []
        while i < len(lines) and '```' not in lines[i]:
            code_lines.append(lines[i])
            i += 1
        if code_lines:
            p = doc.add_paragraph()
            p.style = 'Normal'
            p.add_run('\n'.join(code_lines)).font.name = 'Courier New'
    # 일반 텍스트
    elif line.strip():
        doc.add_paragraph(line)

    i += 1

# DOCX 저장
doc.save(docx_file)
print('Word 파일 생성 완료')
print('파일:', docx_file)
